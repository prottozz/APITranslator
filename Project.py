import os
import yaml
import re
import time
import asyncio
import logging
import json
from datetime import datetime, date, timezone, timedelta
from typing import Dict, List, Optional, Tuple, Set, Any
from pathlib import Path
import shutil
from charset_normalizer import detect
import google.generativeai as genai
from google.generativeai.types import HarmCategory, HarmBlockThreshold
from google.api_core import exceptions as google_exceptions
import aiofiles
import colorlog
# --- НОВЫЙ ИМПОРТ ---
import docx  # Для работы с DOCX
from docx.shared import Pt  # Для указания размера шрифта, если потребуется
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # Для выравнивания
# --- ДОБАВЛЕНЫ ИМПОРТЫ ДЛЯ DOCX MERGE --- (Строка 21)
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.oxml import OxmlElement  # Для добавления разрыва страницы при необходимости

# --- Constants ---
CONFIG_PATH = Path('./config.yml')
CHAPTER_MARKER_TEMPLATE = "---CHAPTER_START_MARKER_ Kapitel {:04d}---"
TRANSLATION_COMPLETE_MARKER = "===TRANSLATION_COMPLETE_MARKER==="
# --- ИЗМЕНЕНИЕ МАРКЕРА ГЛОССАРИЯ (Строка 28 -> 32) ---
GLOSSARY_SEPARATOR = "===GLOSSARY_SECTION_SEPARATOR==="  # Более уникальный маркер
GLOSSARY_FILE_HEADER_TEMPLATE = "--- Глоссарий из главы {:04d} ---\n"
GLOSSARY_FILE_SEPARATOR = "------------------------------\n\n"
DATE_FORMAT = "%Y-%m-%d"
QUOTA_RESET_HOUR_UTC = 7

# --- Setup Logging with Colors ---
# ... (без изменений, строки 36-51 -> 40-55) ...
handler = colorlog.StreamHandler()
handler.setFormatter(colorlog.ColoredFormatter(
    '%(log_color)s[%(asctime)s] [%(levelname)s]: %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S',
    log_colors={
        'DEBUG': 'cyan',
        'INFO': 'blue',
        'WARNING': 'yellow',
        'ERROR': 'red',
        'CRITICAL': 'red,bg_white',
    }
))
logger = colorlog.getLogger(__name__)
logger.addHandler(handler)
logger.setLevel(logging.INFO)
logging.getLogger('asyncio').setLevel(logging.WARNING)
logging.getLogger('google').setLevel(logging.WARNING)
logging.getLogger('google.api_core').setLevel(logging.WARNING)
logging.getLogger('google.auth').setLevel(logging.WARNING)
logging.getLogger('urllib3').setLevel(logging.WARNING)


# --- Configuration Class ---
# ... (без изменений, строки 55-107 -> 59-111) ...
class Config:
    def __init__(self, config_path: Path):
        self.config_path = config_path
        self.data = self._load_config()

    def _load_config(self) -> Dict:
        try:
            with open(self.config_path, 'r', encoding='utf-8') as f:
                return yaml.safe_load(f)
        except FileNotFoundError:
            logger.critical(f"Configuration file not found at: {self.config_path}")
            raise SystemExit(f"Configuration file missing: {self.config_path}")
        except yaml.YAMLError as e:
            logger.critical(f"Error parsing configuration file: {e}")
            raise SystemExit(f"Invalid YAML in config: {e}")

    def get(self, *keys: str, default: Any = None) -> Any:
        value = self.data
        try:
            for key in keys:
                value = value[key]
            return value
        except KeyError:
            # logger.debug(f"Key not found: {'.'.join(keys)}. Returning default: {default}")
            return default
        except TypeError:
            logger.error(
                f"Config structure error: Tried to access key '{keys[-1]}' on non-dictionary element at '{'.'.join(keys[:-1])}'")
            return default

    def set(self, value: Any, *keys: str):
        d = self.data
        try:
            for key in keys[:-1]:
                if key in d and not isinstance(d[key], dict):
                    logger.warning(f"Overwriting non-dict value at config key '{'.'.join(keys[:keys.index(key) + 1])}'")
                    d[key] = {}
                d = d.setdefault(key, {})
            d[keys[-1]] = value
        except TypeError:
            logger.error(
                f"Config structure error: Cannot set value at '{'.'.join(keys)}' because a parent element is not a dictionary.")

    def save(self):
        try:
            api_keys_data = self.get('APIKeys', default={})
            if isinstance(api_keys_data, dict):
                for key_name, key_data in api_keys_data.items():
                    if isinstance(key_data, dict) and 'dateUsedQuota' in key_data and isinstance(
                            key_data['dateUsedQuota'], date):
                        key_data['dateUsedQuota'] = key_data['dateUsedQuota'].strftime(DATE_FORMAT)
            else:
                logger.warning("APIKeys section in config is not a dictionary, cannot format dates.")

            with open(self.config_path, 'w', encoding='utf-8') as f:
                yaml.dump(self.data, f, default_flow_style=False, sort_keys=False, allow_unicode=True)
        except IOError as e:
            logger.error(f"Error writing configuration file: {e}")
        except yaml.YAMLError as e:
            logger.error(f"Error formatting configuration data for saving: {e}")


# --- Helper Functions ---
# ... (get_processed_chapters, update_quota_if_needed, get_available_api_keys без изменений, строки 110-236 -> 114-240) ...
def get_processed_chapters(output_path: Path) -> Set[int]:
    """Scans the output directory and returns a set of processed chapter numbers."""
    processed = set()
    if not output_path.exists():
        output_path.mkdir(parents=True, exist_ok=True)
        return processed
    if not output_path.is_dir():
        logger.error(f"Output path '{output_path}' exists but is not a directory.")
        return processed

    for item in output_path.iterdir():
        if item.is_file() and item.suffix.lower() == '.txt':  # Assumes .txt for processed chapters marker
            match = re.match(r'^(\d{4}).*?', item.name)
            if match:
                try:
                    chapter_num = int(match.group(1))
                    processed.add(chapter_num)
                except ValueError:
                    logger.warning(f"Could not parse chapter number from filename: {item.name}")
    return processed


def get_effective_quota_date_info() -> Tuple[date, date, int]:
    """Возвращает (сегодняшняя_дата_utc, эффективная_дата_для_квоты, текущий_час_utc)."""
    now_utc = datetime.now(timezone.utc)
    today_utc_date = now_utc.date()
    current_utc_hour = now_utc.hour

    # Определяем эффективную дату, для которой сейчас должна использоваться квота
    if current_utc_hour < QUOTA_RESET_HOUR_UTC:
        # До часа сброса мы работаем по квоте предыдущего дня UTC
        effective_date = today_utc_date - timedelta(days=1)
    else:
        # После часа сброса мы работаем по квоте текущего дня UTC
        effective_date = today_utc_date
    return today_utc_date, effective_date, current_utc_hour


def update_quota_if_needed(config: Config):
    """
    Обновляет состояние квоты в конфигурации.
    Сбрасывает usedQuota на 0 и устанавливает dateUsedQuota на сегодняшнюю дату UTC,
    если дата в конфиге старше сегодняшней И текущий час UTC >= QUOTA_RESET_HOUR_UTC.
    Также обеспечивает строковый формат для dateUsedQuota.
    """
    today_utc_date, _, current_utc_hour = get_effective_quota_date_info()

    logger.info(
        f"Running update_quota_if_needed. Today UTC: {today_utc_date.strftime(DATE_FORMAT)}, Current UTC hour: {current_utc_hour}, Reset hour: {QUOTA_RESET_HOUR_UTC}")
    updated_config = False
    api_keys_data = config.get('APIKeys', default={})

    if not isinstance(api_keys_data, dict):
        logger.warning("'APIKeys' section is not a dictionary. Cannot update quotas.")
        return
    if not api_keys_data:
        logger.debug("'APIKeys' section is empty. No quotas to update.")  # Изменено на DEBUG
        return

    for key_name, key_data in api_keys_data.items():
        if not isinstance(key_data, dict):
            logger.warning(f"API key entry '{key_name}' is not a dictionary. Skipping quota update.")
            continue

        account_name = key_data.get('account', key_name)
        original_date_used_quota_from_config = key_data.get('dateUsedQuota')  # Получаем как есть

        try:
            stored_date_obj = None
            if isinstance(original_date_used_quota_from_config, date):
                stored_date_obj = original_date_used_quota_from_config
            elif isinstance(original_date_used_quota_from_config, str):
                try:
                    stored_date_obj = datetime.strptime(original_date_used_quota_from_config, DATE_FORMAT).date()
                except ValueError:
                    logger.error(
                        f"Key '{account_name}': Invalid date string '{original_date_used_quota_from_config}' for 'dateUsedQuota'. Using epoch.")
                    stored_date_obj = date(1970, 1, 1)
            else:  # None или другой тип
                logger.warning(
                    f"Key '{account_name}': 'dateUsedQuota' is missing or has an unexpected type: {type(original_date_used_quota_from_config)}. Assuming very old date (epoch).")
                stored_date_obj = date(1970, 1, 1)  # Если дата отсутствует, считаем её очень старой

            # Основная логика сброса квоты на новый день
            if stored_date_obj < today_utc_date and current_utc_hour >= QUOTA_RESET_HOUR_UTC:
                logger.info(
                    f"Key '{account_name}': Resetting quota for new day. Stored date {stored_date_obj.strftime(DATE_FORMAT)} < Today UTC {today_utc_date.strftime(DATE_FORMAT)} AND current hour {current_utc_hour} >= reset hour {QUOTA_RESET_HOUR_UTC}.")
                config.set(0, 'APIKeys', key_name, 'usedQuota')
                config.set(today_utc_date.strftime(DATE_FORMAT), 'APIKeys', key_name, 'dateUsedQuota')
                updated_config = True
            elif stored_date_obj > today_utc_date:  # Дата в будущем
                logger.warning(
                    f"Key '{account_name}': Stored date {stored_date_obj.strftime(DATE_FORMAT)} is in the future. Check system clocks or config.")
                # Если это был объект date, все равно сохраняем как строку
                if isinstance(original_date_used_quota_from_config, date):
                    config.set(original_date_used_quota_from_config.strftime(DATE_FORMAT), 'APIKeys', key_name,
                               'dateUsedQuota')
                    updated_config = True
            else:  # stored_date_obj == today_utc_date ИЛИ (stored_date_obj < today_utc_date И current_utc_hour < QUOTA_RESET_HOUR_UTC)
                # В этих случаях мы не сбрасываем usedQuota.
                # Просто убедимся, что дата сохранена как строка, если она была объектом date.
                if isinstance(original_date_used_quota_from_config, date):
                    config.set(original_date_used_quota_from_config.strftime(DATE_FORMAT), 'APIKeys', key_name,
                               'dateUsedQuota')
                    updated_config = True
        except Exception as e:
            logger.error(f"Key '{account_name}': Unexpected error during quota update: {e}", exc_info=True)

    if updated_config:
        logger.info("APIKeys configuration potentially updated. Saving config.")
        try:
            config.save()
        except Exception as e:
            logger.error(f"Error saving config after quota update: {e}", exc_info=True)


def get_available_api_keys(config: Config) -> List[Tuple[str, Dict]]:
    available = []
    today_utc_date, effective_quota_date, current_utc_hour = get_effective_quota_date_info()
    effective_quota_date_str = effective_quota_date.strftime(DATE_FORMAT)

    logger.info(
        f"Running get_available_api_keys. Effective quota date: {effective_quota_date_str}, Current UTC hour: {current_utc_hour}, Reset hour: {QUOTA_RESET_HOUR_UTC}")

    api_keys_data = config.get('APIKeys', default={})
    if not isinstance(api_keys_data, dict):
        logger.warning("'APIKeys' section is not a dictionary. Cannot get available keys.")
        return []
    if not api_keys_data:
        logger.debug("'APIKeys' section is empty. No keys to check.")
        return []

    for key_name, key_data in api_keys_data.items():
        if not isinstance(key_data, dict):
            logger.warning(f"API key entry '{key_name}' is not a dictionary. Skipping.")
            continue

        account_name = key_data.get('account', key_name)
        key_value = key_data.get('key')
        quota_limit_cfg = key_data.get('quota')
        used_quota_cfg = key_data.get('usedQuota')

        stored_date_str_from_config = key_data.get('dateUsedQuota', '1970-01-01')  # По умолчанию строка
        if isinstance(stored_date_str_from_config, date):  # Если YAML загрузил как объект date
            stored_date_str_from_config = stored_date_str_from_config.strftime(DATE_FORMAT)
        else:  # Убедимся, что это строка для сравнения
            stored_date_str_from_config = str(stored_date_str_from_config)

        reason_unavailable = ""
        is_key_available = False  # Флаг доступности

        if not key_value:
            reason_unavailable = "API key value is missing."
        else:
            try:
                quota_limit = int(quota_limit_cfg) if quota_limit_cfg is not None else 0
                used_quota = int(used_quota_cfg) if used_quota_cfg is not None else 0
            except (ValueError, TypeError):
                logger.error(
                    f"Key '{account_name}': Invalid quota/usedQuota values. Q: '{quota_limit_cfg}', U: '{used_quota_cfg}'. Assuming key unavailable.")
                quota_limit = 0;
                used_quota = 0  # Делаем ключ недоступным

            if quota_limit <= 0:
                reason_unavailable = f"Quota limit is {quota_limit} (not > 0)."
            else:
                # Сценарий 1: Дата в конфиге совпадает с эффективной датой квоты
                if stored_date_str_from_config == effective_quota_date_str:
                    if used_quota < quota_limit:
                        is_key_available = True
                    else:
                        reason_unavailable = f"Quota reached ({used_quota}/{quota_limit}) for effective date {effective_quota_date_str}."
                # Сценарий 2: Пограничный случай - сейчас >= 7 утра, ожидаем сегодняшнюю квоту,
                # но в конфиге еще вчерашняя дата (update_quota_if_needed еще не сбросила).
                # Считаем, что доступна полная новая квота.
                elif current_utc_hour >= QUOTA_RESET_HOUR_UTC and \
                        effective_quota_date == today_utc_date and \
                        stored_date_str_from_config == (today_utc_date - timedelta(days=1)).strftime(DATE_FORMAT):
                    logger.info(
                        f"Key '{account_name}': Effective date is today ({today_utc_date.strftime(DATE_FORMAT)}), stored date is yesterday. Assuming new day's quota (0/{quota_limit}) is available.")
                    is_key_available = True  # Предполагаем, что usedQuota для этого нового дня будет 0
                # Иначе - даты не совпадают, и это не пограничный случай
                else:
                    reason_unavailable = (f"Stored date '{stored_date_str_from_config}' does not match "
                                          f"effective quota date '{effective_quota_date_str}'.")

        if is_key_available:
            available.append((key_name, key_data))
            logger.debug(
                f"Key '{account_name}' ({key_name}) is available (Effective date: {effective_quota_date_str}).")
        else:
            logger.debug(
                f"Key '{account_name}' ({key_name}) unavailable: {reason_unavailable} (Effective date: {effective_quota_date_str})")

    if not available:
        logger.warning(
            f"No API keys are currently available (Effective quota date: {effective_quota_date_str}, current hour: {current_utc_hour}).")
    else:
        logger.info(f"Found {len(available)} available API key(s) (Effective quota date: {effective_quota_date_str}).")

    return available

# --- ИЗМЕНЕНИЕ: Загрузка только файлов глоссария, начинающихся с "Glossary_" (Строка 239 -> 243) ---
async def load_prompt_and_glossaries(prompt_path: Path, glossary_path: Path) -> str:
    """Loads the main prompt and appends content from glossary files starting with 'Glossary_'."""
    prompt_contents = "Translate the text."
    try:
        async with aiofiles.open(prompt_path, "r", encoding="utf-8") as promptFile:
            prompt_contents = await promptFile.read()
    except FileNotFoundError:
        logger.error(f"Prompt file not found: {prompt_path}. Using default prompt.")
    except Exception as e:
        logger.error(f"Error reading prompt file {prompt_path}: {e}")

    if glossary_path.is_dir():
        logger.info(f"Loading glossaries starting with 'Glossary_' from {glossary_path}...")
        glossary_count = 0
        # --- ИЗМЕНЕНИЕ: Фильтруем файлы по имени (Строка 256 -> 260) ---
        glossary_files = sorted(list(glossary_path.glob('Glossary_*.txt')))  # Ищем только Glossary_*.txt
        if not glossary_files:
            logger.info("No files starting with 'Glossary_' found in glossary directory.")
        for item in glossary_files:
            try:
                async with aiofiles.open(item, "r", encoding="utf-8") as glossary_file:
                    glossary_content = await glossary_file.read()
                    prompt_contents += f"\n\n# Glossary: {item.name}\n{glossary_content}"
                    glossary_count += 1
            except Exception as e:
                logger.warning(f"Could not read glossary file {item}: {e}")
        logger.info(f"Loaded {glossary_count} glossary file(s).")
    else:
        logger.info(f"Glossary path not found or not a directory: {glossary_path}")

    prompt_contents += f"\n\nIMPORTANT: At the very end of the entire translation output, add the exact line:\n{TRANSLATION_COMPLETE_MARKER}"
    return prompt_contents


# --- Core Translation Logic ---
# ... (generate_translation без изменений, строки 274-404 -> 278-408) ...
async def generate_translation(
        prompt: str,
        source_text: str,
        api_key: str,
        config: Config,
        context_info: str = ""
) -> Optional[str]:
    max_retries = config.get('Settings', 'MaxRetries', default=3)
    retry_delay = config.get('Settings', 'RetryDelay', default=5)
    api_call_delay = config.get('Settings', 'ApiCallDelay', default=2)  # Default changed based on typical usage.
    model_name = config.get('Settings', 'ModelName', default="gemini-1.5-pro-latest")  # Updated default model
    request_timeout = config.get('Settings', 'RequestTimeout', default=600)

    await asyncio.sleep(api_call_delay)

    for attempt in range(max_retries + 1):
        logger.debug(f"API call attempt {attempt + 1}/{max_retries + 1} for {context_info} using model {model_name}")
        try:
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel(
                model_name=model_name,
                system_instruction=prompt,
                safety_settings={
                    HarmCategory.HARM_CATEGORY_HARASSMENT: HarmBlockThreshold.BLOCK_NONE,
                    HarmCategory.HARM_CATEGORY_HATE_SPEECH: HarmBlockThreshold.BLOCK_NONE,
                    HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT: HarmBlockThreshold.BLOCK_NONE,
                    HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: HarmBlockThreshold.BLOCK_NONE,
                },
                generation_config=genai.types.GenerationConfig(
                    response_mime_type="text/plain",
                )
            )

            response = await model.generate_content_async(
                contents=source_text,
                request_options={'timeout': request_timeout}
            )

            if not response.candidates:
                block_reason = "Unknown"
                finish_reason = "Unknown"
                safety_ratings = []
                try:
                    if response.prompt_feedback:
                        block_reason = response.prompt_feedback.block_reason.name
                        safety_ratings = response.prompt_feedback.safety_ratings
                    if hasattr(response,
                               'candidates') and response.candidates:  # Should not be needed if not response.candidates is true
                        if response.candidates[0].finish_reason:
                            finish_reason = response.candidates[0].finish_reason.name
                        if hasattr(response.candidates[0], 'safety_ratings') and response.candidates[0].safety_ratings:
                            safety_ratings = response.candidates[0].safety_ratings
                except (AttributeError, IndexError, ValueError) as feedback_err:
                    logger.warning(f"Could not retrieve full feedback/finish reason for {context_info}: {feedback_err}")

                logger.warning(
                    f"No valid candidates received for {context_info}. "
                    f"Finish reason: {finish_reason}. Block reason: {block_reason}. "
                    f"Attempt {attempt + 1}/{max_retries + 1}. "
                    f"Safety Ratings: {safety_ratings}"
                )
                should_retry = True
                if block_reason not in ["BLOCK_REASON_UNSPECIFIED", "UNKNOWN"]:
                    logger.error(
                        f"Content blocked due to safety reasons ({block_reason}) for {context_info}. No retry.")
                    should_retry = False
                elif finish_reason in ["SAFETY", "RECITATION"]:
                    logger.error(f"Content blocked due to ({finish_reason}) for {context_info}. No retry.")
                    should_retry = False
                elif finish_reason == "MAX_TOKENS":
                    logger.error(
                        f"Response stopped due to max tokens limit ({finish_reason}) for {context_info}. Check model limits and input size. No retry.")
                    should_retry = False
                elif finish_reason in ["OTHER", "UNKNOWN", "FINISH_REASON_UNSPECIFIED"]:
                    logger.warning(
                        f"Response stopped due to unspecified reason ({finish_reason}). Retrying might help.")

                if attempt < max_retries and should_retry:
                    await asyncio.sleep(retry_delay * (attempt + 1))
                    continue
                else:
                    logger.error(
                        f"Failed to get content for {context_info} due to blocking or no candidates after retries.")
                    return None
            try:
                output_text = response.text
            except ValueError as text_err:  # This can happen if response has no .text (e.g. blocked)
                finish_reason_name = 'N/A'
                if hasattr(response, 'candidates') and response.candidates and response.candidates[0].finish_reason:
                    finish_reason_name = response.candidates[0].finish_reason.name
                logger.error(
                    f"Error accessing response text for {context_info}: {text_err}. Finish Reason: {finish_reason_name}")
                if attempt < max_retries:
                    await asyncio.sleep(retry_delay)
                    continue
                else:
                    logger.error(f"Failed to access response text after retries for {context_info}.")
                    return None

            if not output_text.strip().endswith(TRANSLATION_COMPLETE_MARKER):
                logger.warning(
                    f"Incomplete response detected (missing marker) for {context_info}. Output length: {len(output_text)}. Attempt {attempt + 1}/{max_retries + 1}.")
                logger.debug(f"Received partial text (last 100 chars): ...{output_text.strip()[-100:]}")
                if attempt < max_retries:
                    await asyncio.sleep(retry_delay)
                    continue
                else:
                    logger.error(
                        f"Failed to get complete response for {context_info} after {max_retries + 1} attempts due to missing marker.")
                    try:
                        error_filename = Path(
                            f"./ERROR_{Path(context_info).stem}_incomplete_{datetime.now():%Y%m%d_%H%M%S}.txt")
                        async with aiofiles.open(error_filename, "w", encoding="utf-8") as err_file:
                            await err_file.write(output_text)
                        logger.info(f"Saved incomplete response to {error_filename}")
                    except Exception as save_err:
                        logger.error(f"Failed to save incomplete response: {save_err}")
                    return None

            output_text = output_text.rsplit(TRANSLATION_COMPLETE_MARKER, 1)[0].strip()
            logger.debug(f"Successfully translated {context_info}. Length: {len(output_text)}")
            return output_text

        except json.JSONDecodeError as e:  # Should not happen with text/plain
            logger.warning(
                f"JSONDecodeError encountered for {context_info}: {e}. Attempt {attempt + 1}/{max_retries + 1}.")
            if attempt < max_retries:
                await asyncio.sleep(retry_delay); continue
            else:
                logger.error(f"Failed JSONDecodeError after retries for {context_info}."); return None
        except google_exceptions.ResourceExhausted as e:
            logger.warning(f"Quota exceeded for API key during call for {context_info}: {e}")
            return "QUOTA_EXCEEDED"  # Special string to indicate quota issue
        except (google_exceptions.InternalServerError, google_exceptions.ServiceUnavailable) as e:
            error_type = type(e).__name__;
            logger.warning(f"{error_type} from API for {context_info}: {e}. Attempt {attempt + 1}/{max_retries + 1}.")
            if attempt < max_retries:
                await asyncio.sleep(retry_delay * (attempt + 1)); continue
            else:
                logger.error(f"Failed persistent {error_type} for {context_info}."); return None
        except google_exceptions.DeadlineExceeded as e:
            logger.warning(
                f"API call timed out for {context_info} after {request_timeout}s: {e}. Attempt {attempt + 1}/{max_retries + 1}.")
            if attempt < max_retries:
                await asyncio.sleep(retry_delay); continue
            else:
                logger.error(f"Failed persistent timeouts for {context_info}."); return None
        except google_exceptions.InvalidArgument as e:
            logger.error(
                f"Invalid argument passed to API for {context_info}: {e}. Check prompt/text/settings. No retry.",
                exc_info=False)
            return None
        except genai.types.BlockedPromptException as e:  # This is for prompt, response blocking is handled by lack of candidates
            logger.error(f"Prompt blocked for {context_info}. Block reason: {e}. No retry.")
            return None
        except Exception as e:
            logger.error(f"Unexpected error during translation for {context_info}: {e}", exc_info=True)
            return None  # General failure

    logger.error(f"Translation definitively failed for {context_info} after all {max_retries + 1} attempts.")
    return None


# --- File Processing Logic ---
async def process_single_file(
        source_file_path: Path,
        output_path: Path,
        api_key_name: str,
        api_key_data: Dict,
        prompt: str,
        config: Config,
        use_last_successful: bool
) -> Tuple[bool, bool]:
    """Translates a single file. Returns (success, quota_exhausted)."""
    filename = source_file_path.name
    output_file_path = output_path / filename
    chapter_num = -1
    try:
        chapter_match = re.match(r'^(\d{4})', filename)
        if chapter_match: chapter_num = int(chapter_match.group(1))
    except ValueError:
        pass

    account_name = api_key_data.get('account', api_key_name)
    logger.info(f"[{account_name}] Translating single file: {filename}")
    source_contents: Optional[str] = None
    default_encoding = config.get('Settings', 'DefaultEncoding', default='utf-8')
    detected_encoding: Optional[str] = None

    try:  # Блок чтения исходного файла
        try:
            with open(source_file_path, 'rb') as f_detect_bytes:
                file_bytes = f_detect_bytes.read()
            if not file_bytes:
                logger.warning(f"File {filename} is empty. Skipping processing.")
                return False, False
            detected_result = detect(file_bytes)
            if detected_result and detected_result['encoding']:
                detected_encoding = detected_result['encoding'].replace('_', '-').lower()
                # confidence = detected_result.get('confidence', 0) # Закомментировано, т.к. не используется
                # logger.debug(f"Detected encoding for {filename}: {detected_encoding} (Confidence: {confidence:.2f})")
            else:
                logger.warning(
                    f"Charset detection failed or returned None for {filename}. Using default encoding: '{default_encoding}'.")
                detected_encoding = default_encoding
        except Exception as enc_e:
            logger.warning(f"Could not detect encoding for {filename}, using default '{default_encoding}': {enc_e}")
            detected_encoding = default_encoding

        # logger.debug(f"Attempting to read {filename} with encoding '{detected_encoding}'") # Можно убрать
        async with aiofiles.open(source_file_path, "r", encoding=detected_encoding, errors='strict') as source_file:
            source_contents = await source_file.read()

    except UnicodeDecodeError as ude:
        logger.critical(
            f"FATAL: UnicodeDecodeError for {filename}. Tried encoding '{detected_encoding}'. Error: {ude}. "
            f"Check the file or the DefaultEncoding ('{default_encoding}') in config. Stopping."
        )
        raise SystemExit(f"Encoding error in file: {filename}")
    except FileNotFoundError:
        logger.error(f"Source file not found: {source_file_path}"); return False, False
    except LookupError:
        logger.critical(
            f"FATAL: Unknown encoding '{detected_encoding}' used for {filename}. "
            f"Check the DefaultEncoding ('{default_encoding}') in config or file content. Stopping."
        )
        raise SystemExit(f"Unknown encoding: {detected_encoding}")
    except Exception as read_e:
        logger.error(f"Error reading file {filename}: {read_e}", exc_info=True); return False, False

    if source_contents is None: logger.error(f"Failed to read content from {filename}. Skipping."); return False, False

    quota_used_this_call = False
    try:  # Блок перевода и сохранения
        current_config_state = Config(config.config_path)
        current_key_state = current_config_state.get('APIKeys', api_key_name, default={})
        current_used = current_key_state.get('usedQuota', 0)
        quota_limit = current_key_state.get('quota', 0)

        if current_used >= quota_limit:
            logger.warning(
                f"Quota already full for {account_name} before API call for {filename}. ({current_used}/{quota_limit})")
            return False, True

        _, effective_quota_date_for_saving, _ = get_effective_quota_date_info()

        new_used_quota = current_used + 1
        config.set(new_used_quota, 'APIKeys', api_key_name, 'usedQuota')
        # Устанавливаем ДАТУ, для которой была использована эта квота
        config.set(effective_quota_date_for_saving.strftime(DATE_FORMAT), 'APIKeys', api_key_name, 'dateUsedQuota')
        config.save()
        quota_used_this_call = True
        logger.debug(
            f"Incremented quota for {account_name} to {new_used_quota}/{quota_limit} for file {filename}. Date set to {effective_quota_date_for_saving.strftime(DATE_FORMAT)}")

        translated_text = await generate_translation(prompt, source_contents, api_key_data['key'], config,
                                                     context_info=filename)

        if translated_text is None:
            logger.error(f"Translation failed permanently for {filename}.")
            return False, False
        elif translated_text == "QUOTA_EXCEEDED":
            logger.warning(f"Quota exceeded for {account_name} during processing of {filename}. Marking key as full.")
            # Убедимся, что дата также актуализируется для этой отметки
            config.set(quota_limit, 'APIKeys', api_key_name, 'usedQuota')
            config.set(effective_quota_date_for_saving.strftime(DATE_FORMAT), 'APIKeys', api_key_name, 'dateUsedQuota')
            config.save()
            return False, True

            # --- НАЧАЛО ИЗМЕНЕНИЯ: Удаление ведущих пустых строк ---
        if translated_text:  # Убедимся, что текст не пустой перед обработкой
            lines = translated_text.splitlines()
            first_non_empty_line_idx = 0
            while first_non_empty_line_idx < len(lines) and not lines[first_non_empty_line_idx].strip():
                first_non_empty_line_idx += 1

            final_text_to_save = "\n".join(lines[first_non_empty_line_idx:])
            if first_non_empty_line_idx > 0:
                logger.debug(
                    f"Removed {first_non_empty_line_idx} leading empty line(s) from translated output for {filename}.")
        else:
            final_text_to_save = ""  # Если translated_text пуст, сохраняем пустую строку
        # --- КОНЕЦ ИЗМЕНЕНИЯ ---

        output_path.mkdir(parents=True, exist_ok=True)
        async with aiofiles.open(output_file_path, "w", encoding="utf-8") as output_file:
            await output_file.write(final_text_to_save)  # Сохраняем обработанный текст
        logger.info(f"Successfully translated and saved: {output_file_path}")

        if use_last_successful and chapter_num != -1:
            config.set(chapter_num, 'State', 'LastSuccessfulChapter')
            # logger.debug(f"Updated LastSuccessfulChapter to {chapter_num}") # Можно убрать
        current_run_count = config.get('State', 'CurrentRunFilesCount', default=0) + 1
        config.set(current_run_count, 'State', 'CurrentRunFilesCount')
        config.save()
        return True, False

    except SystemExit:
        raise
    except Exception as e:
        logger.error(f"Error processing file {filename}: {e}", exc_info=True)
        return False, quota_used_this_call


async def merge_and_process_chunk(
        chapters_to_process: List[Tuple[int, Path]],
        output_path: Path,
        api_key_name: str,
        api_key_data: Dict,
        prompt: str,
        config: Config,
        use_last_successful: bool
) -> Tuple[bool, bool]:
    """Merges files, translates chunk, splits result. Returns (success, quota_exhausted)."""
    if not chapters_to_process: return True, False
    chapters_to_process.sort()
    first_chapter, last_chapter = chapters_to_process[0][0], chapters_to_process[-1][0]
    chunk_info = f"chapters {first_chapter:04d}-{last_chapter:04d}"
    account_name = api_key_data.get('account', api_key_name)
    logger.info(f"[{account_name}] Merging and translating chunk: {chunk_info}")

    merged_content = ""
    default_encoding = config.get('Settings', 'DefaultEncoding', default='utf-8')

    try:  # Блок слияния исходных файлов
        for chapter_num, source_file_path in chapters_to_process:
            marker = CHAPTER_MARKER_TEMPLATE.format(chapter_num) + "\n"
            detected_encoding: Optional[str] = None
            try:
                with open(source_file_path, 'rb') as f_detect_bytes:
                    file_bytes = f_detect_bytes.read()
                if not file_bytes: logger.warning(
                    f"File {source_file_path.name} in chunk is empty. Skipping merge."); continue
                detected_result = detect(file_bytes)
                if detected_result and detected_result['encoding']:
                    detected_encoding = detected_result['encoding'].replace('_', '-').lower()
                    # logger.debug(f"Detected encoding for {source_file_path.name} in chunk: {detected_encoding}")
                else:
                    logger.warning(
                        f"Detection failed/None for {source_file_path.name} (chunk). Using default: '{default_encoding}'.")
                    detected_encoding = default_encoding
            except Exception as enc_e:
                logger.warning(
                    f"Encoding detection failed for {source_file_path.name} (chunk), using default '{default_encoding}': {enc_e}")
                detected_encoding = default_encoding
            try:
                # logger.debug(f"Attempting to read {source_file_path.name} with encoding '{detected_encoding}' for merge")
                async with aiofiles.open(source_file_path, "r", encoding=detected_encoding, errors='strict') as sf:
                    merged_content += marker + await sf.read() + "\n\n"
            except UnicodeDecodeError as ude:
                logger.critical(
                    f"FATAL: Encoding error in {source_file_path.name} of chunk {chunk_info} (tried {detected_encoding}): {ude}"); raise SystemExit(
                    f"Encoding error in file: {source_file_path.name}")
            except FileNotFoundError:
                logger.error(
                    f"File {source_file_path.name} not found during merge for chunk {chunk_info}. Skipping chunk."); return False, False
            except LookupError:
                logger.critical(
                    f"FATAL: Unknown encoding '{detected_encoding}' for {source_file_path.name}. Check DefaultEncoding ('{default_encoding}')."); raise SystemExit(
                    f"Unknown encoding: {detected_encoding}")
            except Exception as read_e:
                logger.error(f"Error reading {source_file_path.name} during merge: {read_e}",
                             exc_info=True); return False, False

        if not merged_content.strip():
            logger.warning(f"Merged content for chunk {chunk_info} is empty. Skipping API call.")
            return False, False

        merged_prompt = (
            f"You are translating a series of book chapters. Chapters are separated by markers like '{CHAPTER_MARKER_TEMPLATE.format(1234)}'.\n"
            f"Translate the content for chapters {first_chapter}-{last_chapter}.\n"
            f"IMPORTANT: Preserve the chapter markers EXACTLY as they appear in the input, each on its own line, before the translated content of that chapter.\n"
            f"Original prompt instructions:\n{prompt}")

        quota_used_this_call = False
        current_config_state = Config(config.config_path)
        current_key_state = current_config_state.get('APIKeys', api_key_name, default={})
        current_used = current_key_state.get('usedQuota', 0)
        quota_limit = current_key_state.get('quota', 0)

        if current_used >= quota_limit:
            logger.warning(
                f"Quota already full for {account_name} before API call for chunk {chunk_info}. ({current_used}/{quota_limit})")
            return False, True

        _, effective_quota_date_for_saving, _ = get_effective_quota_date_info()

        new_used_quota = current_used + 1
        config.set(new_used_quota, 'APIKeys', api_key_name, 'usedQuota')
        config.set(effective_quota_date_for_saving.strftime(DATE_FORMAT), 'APIKeys', api_key_name, 'dateUsedQuota')
        config.save()
        quota_used_this_call = True
        logger.debug(
            f"Incremented quota for {account_name} to {new_used_quota}/{quota_limit} for chunk {chunk_info}. Date set to {effective_quota_date_for_saving.strftime(DATE_FORMAT)}")

        translated_merged_text = await generate_translation(merged_prompt, merged_content, api_key_data['key'], config,
                                                            context_info=chunk_info)

        if translated_merged_text is None:
            logger.error(f"Translation failed permanently for chunk {chunk_info}.")
            return False, False
        elif translated_merged_text == "QUOTA_EXCEEDED":
            logger.warning(f"Quota exceeded for {account_name} processing chunk {chunk_info}. Marking key as full.")
            config.set(quota_limit, 'APIKeys', api_key_name, 'usedQuota')
            config.set(effective_quota_date_for_saving.strftime(DATE_FORMAT), 'APIKeys', api_key_name, 'dateUsedQuota')
            config.save()
            return False, True

        output_path.mkdir(parents=True, exist_ok=True)
        processed_count_in_chunk = 0
        max_successfully_saved_chapter_in_chunk = config.get('State', 'LastSuccessfulChapter', default=0)

        marker_find_pattern = re.compile(
            r"^" + re.escape(CHAPTER_MARKER_TEMPLATE.split('{:04d}')[0]) + r"(\d{4})" + re.escape(
                CHAPTER_MARKER_TEMPLATE.split('{:04d}')[1]) + r"$", re.MULTILINE)
        matches = list(marker_find_pattern.finditer(translated_merged_text))

        if not matches:
            logger.error(f"No chapter markers found in the translated output for chunk {chunk_info}. Cannot split.")
            error_filename = output_path / f"ERROR_CHUNK_{first_chapter:04d}-{last_chapter:04d}_no_markers_{datetime.now():%Y%m%d_%H%M%S}.txt"
            try:
                async with aiofiles.open(error_filename, "w", encoding="utf-8") as err_file:
                    await err_file.write(translated_merged_text)
                logger.info(f"Saved full response with errors to {error_filename}")
            except Exception as save_e:
                logger.error(f"Failed to save error response: {save_e}")
            return False, False

        logger.info(f"Found {len(matches)} chapter markers in translated output for chunk {chunk_info}.")
        original_chapter_numbers_in_chunk = {num for num, _ in chapters_to_process}

        for i, match in enumerate(matches):
            try:
                chapter_num_split = int(match.group(1))
            except (IndexError, ValueError):
                logger.warning(
                    f"Could not parse chapter number from marker '{match.group(0)}' in chunk {chunk_info}. Skipping this part.")
                continue

            content_start_pos = match.end()
            content_end_pos = matches[i + 1].start() if (i + 1) < len(matches) else len(translated_merged_text)

            # Исходный текст главы (может содержать ведущие/конечные \n от API или разделения)
            raw_content_part = translated_merged_text[content_start_pos:content_end_pos]

            # --- НАЧАЛО ИЗМЕНЕНИЯ: Удаление ведущих пустых строк из content_part ---
            if raw_content_part:  # Проверяем, что есть что обрабатывать
                lines = raw_content_part.splitlines()
                first_non_empty_line_idx = 0
                # Пропускаем пустые строки и строки, состоящие только из пробелов, в начале
                while first_non_empty_line_idx < len(lines) and not lines[first_non_empty_line_idx].strip():
                    first_non_empty_line_idx += 1

                # Собираем текст обратно, начиная с первой непустой строки, сохраняя структуру
                content_part_cleaned = "\n".join(lines[first_non_empty_line_idx:])

                # Также уберем конечные пустые строки/пробелы, которые могли остаться после splitlines().join()
                # или если исходный raw_content_part заканчивался на \n\n
                content_part_final = content_part_cleaned.strip()

                if first_non_empty_line_idx > 0 and raw_content_part.strip():  # Логируем только если были удалены строки и был непустой контент
                    logger.debug(
                        f"Removed leading empty/whitespace lines from chapter {chapter_num_split} content in chunk {chunk_info}.")
            else:
                content_part_final = ""
            # --- КОНЕЦ ИЗМЕНЕНИЯ ---

            if chapter_num_split not in original_chapter_numbers_in_chunk:
                logger.warning(
                    f"Marker found for chapter {chapter_num_split} in chunk {chunk_info}, but it was not in the original list. Skipping save.")
                continue
            if not content_part_final:  # Проверяем уже очищенный и стрипнутый контент
                logger.warning(
                    f"Found marker for chapter {chapter_num_split} in chunk {chunk_info}, but extracted content is empty after cleaning. Skipping save.")
                continue

            output_filename = f"{chapter_num_split:04d}.txt"
            output_file_path = output_path / output_filename
            try:
                async with aiofiles.open(output_file_path, "w", encoding="utf-8") as of:
                    await of.write(content_part_final)  # Сохраняем окончательно очищенный текст
                logger.info(f"Successfully extracted and saved: {output_file_path} from chunk {chunk_info}")
                processed_count_in_chunk += 1
                if chapter_num_split > max_successfully_saved_chapter_in_chunk:
                    max_successfully_saved_chapter_in_chunk = chapter_num_split
            except Exception as write_e:
                logger.error(f"Error writing split file {output_filename} from chunk {chunk_info}: {write_e}")

        if processed_count_in_chunk > 0:
            if use_last_successful:
                current_last_successful = config.get('State', 'LastSuccessfulChapter', default=0)
                if max_successfully_saved_chapter_in_chunk > current_last_successful:
                    config.set(max_successfully_saved_chapter_in_chunk, 'State', 'LastSuccessfulChapter')
                    # logger.debug(f"Updated LastSuccessfulChapter to {max_successfully_saved_chapter_in_chunk} after chunk {chunk_info}") # Можно убрать
            current_run_count = config.get('State', 'CurrentRunFilesCount', default=0) + processed_count_in_chunk
            config.set(current_run_count, 'State', 'CurrentRunFilesCount')
            config.save()

        if processed_count_in_chunk != len(chapters_to_process):
            logger.warning(
                f"Mismatch in processed chapters for chunk {chunk_info}. "
                f"Expected {len(chapters_to_process)}, actually saved {processed_count_in_chunk}."
            )
        return True, False

    except SystemExit:
        raise
    except Exception as e:
        logger.error(f"General error processing chunk {chunk_info}: {e}", exc_info=True)
        return False, quota_used_this_call


    # --- Orchestrators ---
# ... (main_async, main_sequential без изменений, строки 666-863 -> 670-867) ...
# --- START OF MODIFIED FILE Project.py ---
# ... (весь предыдущий код до функции main_async) ...

async def main_async(config: Config):
    """Main asynchronous execution flow."""
    source_path = Path(config.get('Settings', 'SourcePath', default='./Source'))
    output_path = Path(config.get('Settings', 'OutputPath', default='./Output'))
    prompt_path = Path(config.get('Settings', 'PromptPath', default='prompt.txt'))
    glossary_path = Path(config.get('Settings', 'GlossaryPath', default='./Glossaries'))
    end_chapter = config.get('Settings', 'EndChapter', default=10000)
    files_per_run = config.get('Settings', 'FilesPerRun', default=-1)
    merge_chunk_size = config.get('Settings', 'MergeChunkSize', default=0)  # For translation
    use_last_successful = config.get('Settings', 'UseLastSuccessfulChapter', default=True)

    logger.info(
        f"Starting async run. UseLastSuccessfulChapter: {use_last_successful}, MergeChunkSize (translation): {merge_chunk_size}")
    update_quota_if_needed(config)
    available_keys_initial = get_available_api_keys(config)

    if not available_keys_initial:
        logger.warning("No available API keys with quota remaining and reset time passed. Exiting.")
        return

    prompt = await load_prompt_and_glossaries(prompt_path, glossary_path)
    last_successful_chapter = 0
    if use_last_successful:
        last_successful_chapter = config.get('State', 'LastSuccessfulChapter', default=0)
        logger.info(f"Starting from LastSuccessfulChapter: {last_successful_chapter}")
    else:
        logger.info("Starting from chapter 0 (UseLastSuccessfulChapter is false).")
        config.set(0, 'State', 'LastSuccessfulChapter')  # Reset if not using
        config.save()

    processed_in_output = get_processed_chapters(output_path)
    logger.info(f"Found {len(processed_in_output)} chapters already present in {output_path} (used for skipping).")

    files_to_process: List[Tuple[int, Path]] = []
    if source_path.is_dir():
        all_source_files = sorted(list(source_path.glob('[0-9][0-9][0-9][0-9]*.txt')))
        logger.info(f"Found {len(all_source_files)} potential source files in {source_path}.")
        for file_path in all_source_files:
            match = re.match(r'^(\d{4})', file_path.name)
            if match:
                try:
                    chapter_num = int(match.group(1))
                    if chapter_num > last_successful_chapter and \
                            chapter_num <= end_chapter and \
                            chapter_num not in processed_in_output:
                        files_to_process.append((chapter_num, file_path))
                except ValueError:
                    logger.warning(f"Could not parse chapter number from {file_path.name}")
                    continue
    else:
        logger.error(f"Source path '{source_path}' not found or not a directory.")
        return

    if not files_to_process:
        logger.info(
            "No new chapters to process based on current filters (LastSuccessfulChapter, EndChapter, ProcessedInOutput).")
        return
    logger.info(f"Found {len(files_to_process)} chapters to process in this run.")

    actual_files_to_process_this_run = files_to_process
    if files_per_run > 0 and len(files_to_process) > files_per_run:
        actual_files_to_process_this_run = files_to_process[:files_per_run]
        logger.info(
            f"Limiting processing to {len(actual_files_to_process_this_run)} chapters due to FilesPerRun={files_per_run}.")

    if not actual_files_to_process_this_run:
        logger.info("No files left to process after FilesPerRun limit.")
        return

    exhausted_keys = set()
    active_keys = list(available_keys_initial)
    key_cycle = asyncio.Queue()
    for key_info in active_keys:
        await key_cycle.put(key_info)

    semaphore = asyncio.Semaphore(len(active_keys))
    processed_tasks_count = 0

    async def worker(item_to_process: Any, is_chunk: bool):
        nonlocal processed_tasks_count, active_keys, exhausted_keys  # processed_tasks_count не используется, можно убрать nonlocal для него если так

        # Не будем проверять key_cycle.empty() здесь, доверимся семафору
        # и проверке внутри семафора

        async with semaphore:  # Захватываем семафор ПЕРЕД попыткой получить ключ
            if key_cycle.empty():
                # Это более легитимная ситуация: семафор разрешил, но ключей физически нет
                # (все ключи изъяты другими рабочими потоками и, возможно, исчерпаны)
                logger.warning(
                    "Worker (in semaphore): No keys available in cycle at this moment. Task cannot run with this attempt.")
                return False  # Не удалось получить ключ

            api_key_name, api_key_data = await key_cycle.get()
            account_name = api_key_data.get('account', api_key_name)

            # Проверяем, не исчерпан ли ключ уже (могло случиться, пока он был в очереди)
            # Эта проверка важна, так как состояние ключа могло измениться в config.yml
            # другим worker-ом.
            # Однако, если мы полностью доверяем exhausted_keys, то эта проверка может быть избыточной.
            # Оставим exhausted_keys как основной механизм отслеживания.
            if api_key_name in exhausted_keys:
                logger.debug(
                    f"Key {account_name} ({api_key_name}) is in exhausted_keys set. Returning to cycle (will be skipped again).")
                # Возвращаем ключ в очередь, чтобы другие worker-ы, если они есть,
                # также могли его увидеть и пропустить.
                # Или, если ключ точно исчерпан, его можно вообще не возвращать,
                # но тогда key_cycle может стать пустой, и семафор будет блокировать вечно, если нет других ключей.
                # Лучше вернуть, чтобы цикл не завис, если есть активные задачи.
                await key_cycle.put((api_key_name, api_key_data))
                return False  # Задача не выполнена этим worker-ом с этим ключом

            success, quota_exhausted_by_this_call = False, False
            try:
                # Логика выполнения задачи (process_single_file или merge_and_process_chunk)
                if is_chunk:
                    success, quota_exhausted_by_this_call = await merge_and_process_chunk(
                        item_to_process, output_path, api_key_name, api_key_data, prompt, config, use_last_successful
                    )
                else:
                    success, quota_exhausted_by_this_call = await process_single_file(
                        item_to_process[1], output_path, api_key_name, api_key_data, prompt, config, use_last_successful
                    )

                # if success: processed_tasks_count += 1 # Если нужно считать успешно выполненные задачи

            except SystemExit as e:
                logger.critical(f"SystemExit in worker for key {account_name}: {e}. Re-raising.")
                # Важно! Если произошел SystemExit, ключ может быть не возвращен в очередь.
                # Это нормально, так как вся программа завершается.
                raise
            except Exception as e:
                logger.error(f"Unhandled exception in worker with key {account_name} ({api_key_name}): {e}",
                             exc_info=True)
                success = False
                # Попытка определить, исчерпана ли квота, если произошла ошибка
                # Лучше, чтобы process_single_file/merge_and_process_chunk сами возвращали это.
                # Здесь это как запасной вариант.
                cfg_check = Config(config.config_path)  # Свежее чтение конфига
                key_state_after_call = cfg_check.get('APIKeys', api_key_name, default={})
                quota_limit_check = key_state_after_call.get('quota', 0)
                used_quota_check = key_state_after_call.get('usedQuota', 0)
                if quota_limit_check > 0:  # Проверяем, чтобы избежать деления на ноль или некорректной логики
                    quota_exhausted_by_this_call = used_quota_check >= quota_limit_check
            finally:
                if quota_exhausted_by_this_call:
                    logger.warning(
                        f"API key {account_name} ({api_key_name}) was exhausted by this call or found exhausted. Adding to exhausted_keys set.")
                    exhausted_keys.add(api_key_name)
                    # Не возвращаем исчерпанный ключ в очередь key_cycle,
                    # чтобы он не выбирался снова для выполнения задач.
                    # Семафор будет освобожден, но этот ключ больше не будет циркулировать.
                elif api_key_name not in exhausted_keys:  # Если ключ не исчерпан
                    logger.debug(f"Returning key {account_name} ({api_key_name}) to key_cycle.")
                    await key_cycle.put((api_key_name, api_key_data))
                else:
                    # Ключ был в exhausted_keys изначально, и мы его не использовали.
                    # Он уже был возвращен в key_cycle ранее в блоке if api_key_name in exhausted_keys.
                    # Или он только что был добавлен в exhausted_keys, и мы его не возвращаем.
                    logger.debug(
                        f"Key {account_name} ({api_key_name}) is in exhausted_keys set and was not used or just marked. Not returning to cycle again from here.")

                # Семафор освобождается автоматически при выходе из блока 'async with semaphore'
            return success

    items_for_tasks = []
    if merge_chunk_size > 1:
        for i in range(0, len(actual_files_to_process_this_run), merge_chunk_size):
            chunk = actual_files_to_process_this_run[i:i + merge_chunk_size]
            if chunk: items_for_tasks.append((chunk, True))
    else:
        items_for_tasks = [((num, path), False) for num, path in actual_files_to_process_this_run]

    if not items_for_tasks:
        logger.info("No items (single files or chunks) prepared for tasks. Exiting async run.")
        return

    logger.info(f"Preparing to run {len(items_for_tasks)} processing tasks (single files or chunks)...")
    tasks = [asyncio.create_task(worker(item_data, is_chunk_task)) for item_data, is_chunk_task in items_for_tasks]

    try:
        results = await asyncio.gather(*tasks, return_exceptions=True)

        successful_tasks_count = 0
        failed_tasks_count = 0
        for i, res_or_exc in enumerate(results):
            # --- ИЗМЕНЕННЫЙ БЛОК ДЛЯ context_info ---
            # Получаем item_data и is_chunk_task, которые соответствуют текущему результату
            item_data_for_log, is_chunk_for_log = items_for_tasks[i]
            context_info = ""
            if is_chunk_for_log:
                # item_data_for_log это список кортежей [(chapter_num, path), ...]
                if item_data_for_log:  # Проверка, что список не пуст
                    first_chap_num_log = item_data_for_log[0][0]
                    last_chap_num_log = item_data_for_log[-1][0]
                    context_info = f"Chunk {first_chap_num_log:04d}-{last_chap_num_log:04d}"
                else:
                    context_info = "Empty Chunk"  # На случай, если пустой чанк как-то попал
            else:
                # item_data_for_log это кортеж (chapter_num, path)
                chap_num_log = item_data_for_log[0]
                file_name_log = item_data_for_log[1].name
                context_info = f"File {file_name_log} (Chapter {chap_num_log:04d})"
            # --- КОНЕЦ ИЗМЕНЕННОГО БЛОКА ---

            if isinstance(res_or_exc, Exception):
                failed_tasks_count += 1
                if isinstance(res_or_exc, SystemExit):
                    logger.critical(f"SystemExit encountered in a worker task for {context_info}. Stopping run.")
                    raise res_or_exc
                logger.error(f"Task for {context_info} failed with exception: {res_or_exc}",
                             exc_info=(isinstance(res_or_exc, Exception) and res_or_exc or None))
            elif res_or_exc is True:
                successful_tasks_count += 1
            else:
                failed_tasks_count += 1
                logger.warning(f"Task for {context_info} reported failure (returned False).")

        logger.info(
            f"Async run finished. Total tasks: {len(tasks)}. Successful tasks: {successful_tasks_count}, Failed tasks: {failed_tasks_count}")

        final_available_keys = get_available_api_keys(config)
        if not final_available_keys and (failed_tasks_count > 0 or successful_tasks_count < len(tasks)):
            logger.warning(f"Run finished, and all API keys appear to be exhausted or unavailable.")
        elif failed_tasks_count > 0:
            logger.warning(f"{failed_tasks_count} tasks may have failed or were not processed fully. Check logs.")

    except SystemExit:
        logger.critical("Async run terminated due to SystemExit (e.g., encoding error in a file).")
    except asyncio.CancelledError:
        logger.warning("Async run was cancelled.")
    except Exception as main_e:
        logger.critical(f"Critical error during task execution orchestration in main_async: {main_e}", exc_info=True)


async def main_sequential(config: Config):  # Убедимся, что она async
    """Main sequential execution flow."""
    source_path = Path(config.get('Settings', 'SourcePath', default='./Source'))
    output_path = Path(config.get('Settings', 'OutputPath', default='./Output'))
    prompt_path = Path(config.get('Settings', 'PromptPath', default='prompt.txt'))
    glossary_path = Path(config.get('Settings', 'GlossaryPath', default='./Glossaries'))
    end_chapter = config.get('Settings', 'EndChapter', default=10000)
    files_per_run = config.get('Settings', 'FilesPerRun', default=-1)
    api_call_delay = config.get('Settings', 'ApiCallDelay', default=2)
    use_last_successful = config.get('Settings', 'UseLastSuccessfulChapter', default=True)

    logger.info(f"Starting sequential run. UseLastSuccessfulChapter: {use_last_successful}")
    update_quota_if_needed(config)

    prompt = ""
    try:
        # Используем асинхронную функцию загрузки, теперь с await, т.к. main_sequential - async
        prompt = await load_prompt_and_glossaries(prompt_path, glossary_path)
    except Exception as e:
        logger.error(f"Error loading prompt/glossaries: {e}")
        # Убедимся, что TRANSLATION_COMPLETE_MARKER используется из константы
        prompt = f"Translate the text.\n\nIMPORTANT: At the very end of the entire translation output, add the exact line:\n{TRANSLATION_COMPLETE_MARKER}"

    processed_in_output = get_processed_chapters(output_path)
    last_successful_chapter = 0
    if use_last_successful:
        last_successful_chapter = config.get('State', 'LastSuccessfulChapter', default=0)
        logger.info(f"Starting from LastSuccessfulChapter: {last_successful_chapter}")
    else:
        logger.info("Starting from chapter 0 (UseLastSuccessfulChapter is false).")
        config.set(0, 'State', 'LastSuccessfulChapter')
        config.save()

    files_to_process_seq: List[Tuple[int, Path]] = []
    if source_path.is_dir():
        all_source_files = sorted(list(source_path.glob('[0-9][0-9][0-9][0-9]*.txt')))
        for file_path in all_source_files:
            match = re.match(r'^(\d{4})', file_path.name)
            if match:
                try:
                    chapter_num = int(match.group(1))
                    if chapter_num > last_successful_chapter and \
                            chapter_num <= end_chapter and \
                            chapter_num not in processed_in_output:
                        files_to_process_seq.append((chapter_num, file_path))
                except ValueError:
                    continue
    else:
        logger.error(f"Source path '{source_path}' not found."); return

    if not files_to_process_seq:
        logger.info("No new chapters to process sequentially based on current filters.")
        return

    logger.info(f"Found {len(files_to_process_seq)} potential chapters for sequential processing.")
    actual_files_to_process_seq = files_to_process_seq
    if 0 < files_per_run < len(files_to_process_seq):
        actual_files_to_process_seq = files_to_process_seq[:files_per_run]
        logger.info(f"Limiting to {len(actual_files_to_process_seq)} chapters due to FilesPerRun.")

    if not actual_files_to_process_seq:
        logger.info("No files left for sequential processing after FilesPerRun limit.")
        return

    files_processed_count = 0
    all_keys_exhausted_for_run = False

    for chapter_num, file_path in actual_files_to_process_seq:
        if all_keys_exhausted_for_run:
            logger.info(f"Skipping remaining chapters as all keys exhausted during this run.")
            break

        logger.info(f"Attempting chapter {chapter_num} ({file_path.name})...");
        processed_successfully_this_chapter = False

        current_available_keys = get_available_api_keys(config)
        if not current_available_keys:
            logger.warning("No available API keys left for sequential run. Stopping.")
            all_keys_exhausted_for_run = True
            break

        for key_name, key_data in current_available_keys:
            account_name = key_data.get('account', key_name)
            logger.debug(f"Trying key {account_name} for chapter {chapter_num}.")

            success_this_key, quota_exhausted_this_key = False, False
            try:
                await asyncio.sleep(api_call_delay)
                success_this_key, quota_exhausted_this_key = await process_single_file(  # process_single_file уже async
                    file_path, output_path, key_name, key_data, prompt, config, use_last_successful
                )
            except SystemExit as e:
                logger.critical(f"SystemExit during sequential processing of {file_path.name}: {e}")
                raise
            except Exception as e:
                logger.error(f"Error running async process_single_file for {file_path.name} with {account_name}: {e}",
                             exc_info=True)
                success_this_key = False
                cfg_check = Config(config.config_path)
                key_state_after_call = cfg_check.get('APIKeys', key_name, default={})
                quota_limit_check = key_state_after_call.get('quota', 0)
                used_quota_check = key_state_after_call.get('usedQuota', 0)
                if quota_limit_check > 0:
                    quota_exhausted_this_key = used_quota_check >= quota_limit_check
                else:  # Если квота 0 или не задана, считаем, что не исчерпана по этой причине
                    quota_exhausted_this_key = False

            if success_this_key:
                logger.info(f"Chapter {chapter_num} processed successfully with {account_name}.")
                processed_successfully_this_chapter = True
                files_processed_count += 1
                break
            elif quota_exhausted_this_key:
                logger.warning(f"Key {account_name} exhausted on chapter {chapter_num}. Trying next available key.")
                continue
            else:
                logger.error(
                    f"Failed chapter {chapter_num} with {account_name} (non-quota API error or other issue). Trying next available key.")
                continue

        if not processed_successfully_this_chapter:
            logger.error(
                f"Could not process chapter {chapter_num}. All tried keys failed or no keys were suitable for it.")

    logger.info(f"Sequential run finished. Total chapters processed in this run: {files_processed_count}.")
    if all_keys_exhausted_for_run:
        remaining_chapters = len(actual_files_to_process_seq) - files_processed_count
        if remaining_chapters > 0:
            logger.warning(
                f"Run finished because all API keys exhausted or became unavailable, {remaining_chapters} chapters from this run's list were left unprocessed.")


# --- ИЗМЕНЕНИЕ: Функция очистки и извлечения глоссария (Строка 867 -> 871) ---
#   - Используется новый GLOSSARY_SEPARATOR
#   - Удаление начальных пустых строк теперь происходит *перед* записью в CleanedOutput
async def extract_glossary_and_clean_files(config: Config):
    """Extracts glossaries, cleans files (incl. leading empty lines), and saves results, splitting glossaries if configured."""
    output_path = Path(config.get('Settings', 'OutputPath', default='./Output'))
    glossary_path = Path(config.get('Settings', 'GlossaryPath', default='./Glossaries'))
    temp_cleaned_path = Path(
        config.get('Settings', 'TempCleanedPath', default='./TempCleaned'))  # Используется для временных файлов
    cleaned_output_path = Path(
        config.get('Settings', 'CleanedOutputPath', default='./CleanedOutput'))  # Финальный путь для очищенных
    glossary_chapters_per_file = config.get('Settings', 'GlossaryChaptersPerFile',
                                            default=0)  # 0 or less means one file

    logger.info(f"Starting glossary extraction from '{output_path}'.")
    logger.info(f"Glossaries will be saved to '{glossary_path}'.")
    logger.info(f"Cleaned files will be saved to '{cleaned_output_path}'.")
    if glossary_chapters_per_file > 0:
        logger.info(f"Glossaries will be split into files with max {glossary_chapters_per_file} chapters each.")
    else:
        logger.info(
            "All extracted glossaries will be saved into a single file (or multiple if names collide due to ranges).")

    if not output_path.is_dir():
        logger.error(f"Source path for extraction '{output_path}' not found. Cannot proceed.")
        return

    # Создаем/очищаем директории
    glossary_path.mkdir(parents=True, exist_ok=True)
    cleaned_output_path.mkdir(parents=True, exist_ok=True)  # Финальная папка очищенных
    temp_cleaned_path.mkdir(parents=True, exist_ok=True)  # Временная папка для очистки

    # Очистка временной папки перед использованием
    for item in temp_cleaned_path.iterdir():
        try:
            if item.is_file():
                item.unlink()
            elif item.is_dir():
                shutil.rmtree(item)
        except Exception as e:
            logger.warning(f"Could not clear item {item} from temp directory: {e}")

    extracted_glossaries: Dict[int, str] = {}  # {chapter_num: glossary_text}
    files_processed_for_cleaning = 0
    files_with_glossary_found = 0

    # Собираем файлы для обработки
    # Предполагаем, что в OutputPath лежат .txt файлы с номерами глав
    files_in_output = sorted(list(output_path.glob('[0-9][0-9][0-9][0-9]*.txt')))
    if not files_in_output:
        logger.info(f"No files found in {output_path} to process for glossary extraction and cleaning.")
        return
    logger.info(f"Found {len(files_in_output)} files in {output_path} for glossary extraction and cleaning.")

    async def process_file_for_glossary_and_cleaning(file_path: Path):
        nonlocal files_processed_for_cleaning, files_with_glossary_found  # Allow modification of counters
        chapter_num = -1
        # Временный путь для очищенного файла ДО перемещения в CleanedOutputPath
        temp_target_cleaned_file_path = temp_cleaned_path / file_path.name

        try:
            match = re.match(r'^(\d{4})', file_path.name)
            if match:
                chapter_num = int(match.group(1))
            else:
                logger.warning(
                    f"Could not parse chapter number from filename: {file_path.name}. It will be copied as is to temp for cleaning pass.")
                # Копируем как есть во временную папку, если нет номера главы, для последующего общего шага очистки
                # Но глоссарий извлечь не получится.
                try:
                    await asyncio.to_thread(shutil.copy2, file_path, temp_target_cleaned_file_path)
                except Exception as copy_e:
                    logger.error(f"Failed to copy {file_path.name} to temp for cleaning: {copy_e}")
                return  # Не можем извлечь глоссарий

            async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='replace') as infile:
                content = await infile.read()

            cleaned_text_content = content  # По умолчанию весь контент, если нет разделителя
            glossary_text_content = ""

            separator_pos = content.find(GLOSSARY_SEPARATOR)

            if separator_pos != -1:
                cleaned_text_content = content[
                                       :separator_pos].rstrip()  # Текст до разделителя (с удалением пробелов справа)
                glossary_text_content = content[
                                        separator_pos + len(GLOSSARY_SEPARATOR):].strip()  # Текст после разделителя

                if glossary_text_content:  # Если глоссарий не пустой
                    if chapter_num != -1:  # Убедимся, что номер главы есть
                        extracted_glossaries[chapter_num] = glossary_text_content
                        files_with_glossary_found += 1
                        logger.debug(f"Extracted glossary from chapter {chapter_num} ({file_path.name}).")
                else:
                    logger.debug(
                        f"Glossary separator found in chapter {chapter_num} ({file_path.name}), but glossary section is empty.")
            else:
                logger.debug(
                    f"No glossary separator ('{GLOSSARY_SEPARATOR}') found in {file_path.name}. Entire content treated as main text for cleaning.")
                cleaned_text_content = content.strip()  # Убираем пробелы с обоих концов, если разделителя нет

            # --- Очистка от ведущих пустых строк ---
            lines = cleaned_text_content.splitlines()
            first_non_empty_line_idx = 0
            while first_non_empty_line_idx < len(lines) and not lines[first_non_empty_line_idx].strip():
                first_non_empty_line_idx += 1

            # Собираем текст обратно, начиная с первой непустой строки, сохраняя оригинальные переносы строк
            final_cleaned_text_for_file = "\n".join(lines[first_non_empty_line_idx:])

            # Записываем очищенный текст (без глоссария, без ведущих пустых строк) во временный файл
            async with aiofiles.open(temp_target_cleaned_file_path, 'w', encoding='utf-8') as outfile:
                await outfile.write(final_cleaned_text_for_file)
            logger.debug(f"Saved cleaned content for {file_path.name} to temp path {temp_target_cleaned_file_path}.")
            files_processed_for_cleaning += 1

        except Exception as e:
            logger.error(f"Error processing file {file_path.name} for glossary/cleaning: {e}", exc_info=True)
            # Попытка скопировать исходный файл во временную папку в случае ошибки, чтобы он не потерялся
            try:
                await asyncio.to_thread(shutil.copy2, file_path, temp_target_cleaned_file_path)
                logger.warning(f"Copied original file {file_path.name} to temp path due to processing error.")
            except Exception as copy_e:
                logger.error(f"Failed to copy original file {file_path.name} to temp path after error: {copy_e}")

    # Запускаем задачи для обработки файлов
    tasks = [asyncio.create_task(process_file_for_glossary_and_cleaning(fp)) for fp in files_in_output]
    if tasks:
        await asyncio.gather(*tasks)

    logger.info(f"File processing for glossary extraction and temp cleaning complete. "
                f"Files processed for cleaning: {files_processed_for_cleaning}, "
                f"Files with glossary found: {files_with_glossary_found}.")

    # --- Сохранение извлеченных глоссариев ---
    if extracted_glossaries:
        sorted_chapter_numbers_with_glossaries = sorted(extracted_glossaries.keys())
        num_glossaries_to_save = len(sorted_chapter_numbers_with_glossaries)
        logger.info(f"Found {num_glossaries_to_save} glossaries to save.")

        if glossary_chapters_per_file <= 0:  # Сохранить все в один файл
            if sorted_chapter_numbers_with_glossaries:  # Если есть что сохранять
                min_chap = sorted_chapter_numbers_with_glossaries[0]
                max_chap = sorted_chapter_numbers_with_glossaries[-1]
                # Имя файла для глоссария теперь включает префикс "Glossary_"
                glossary_filename = glossary_path / f"{min_chap:04d}-{max_chap:04d}.txt"
                logger.info(
                    f"Saving combined glossary ({num_glossaries_to_save} chapters: {min_chap:04d}-{max_chap:04d}) to {glossary_filename}...")
                try:
                    async with aiofiles.open(glossary_filename, 'w', encoding='utf-8') as f_glossary:
                        for chapter_num in sorted_chapter_numbers_with_glossaries:
                            await f_glossary.write(GLOSSARY_FILE_HEADER_TEMPLATE.format(chapter_num))
                            await f_glossary.write(
                                extracted_glossaries[chapter_num] + "\n")  # Добавляем \n после текста глоссария
                            await f_glossary.write(
                                GLOSSARY_FILE_SEPARATOR)  # Добавляем разделитель между глоссариями глав
                    logger.info(f"Combined glossary saved successfully to {glossary_filename}.")
                except Exception as e:
                    logger.error(f"Error saving combined glossary file {glossary_filename}: {e}")
        else:  # Разделить глоссарии на несколько файлов
            logger.info(
                f"Splitting {num_glossaries_to_save} glossaries into files with max {glossary_chapters_per_file} chapters each.")
            saved_glossary_chunks_count = 0
            for i in range(0, num_glossaries_to_save, glossary_chapters_per_file):
                chunk_of_chapter_numbers = sorted_chapter_numbers_with_glossaries[i: i + glossary_chapters_per_file]
                if not chunk_of_chapter_numbers: continue

                chunk_min_chap = chunk_of_chapter_numbers[0]
                chunk_max_chap = chunk_of_chapter_numbers[-1]
                # Имя файла для чанка глоссария
                # Убедимся, что имя файла для чанка также начинается с "Glossary_"
                chunk_filename = glossary_path / f"{chunk_min_chap:04d}-{chunk_max_chap:04d}.txt"
                logger.info(
                    f"Saving glossary chunk ({len(chunk_of_chapter_numbers)} chapters: {chunk_min_chap:04d}-{chunk_max_chap:04d}) to {chunk_filename}...")
                try:
                    async with aiofiles.open(chunk_filename, 'w', encoding='utf-8') as f_chunk_glossary:
                        for chapter_num in chunk_of_chapter_numbers:
                            if chapter_num in extracted_glossaries:  # Проверка на всякий случай
                                await f_chunk_glossary.write(GLOSSARY_FILE_HEADER_TEMPLATE.format(chapter_num))
                                await f_chunk_glossary.write(extracted_glossaries[chapter_num] + "\n")
                                await f_chunk_glossary.write(GLOSSARY_FILE_SEPARATOR)
                    logger.info(f"Glossary chunk {chunk_filename.name} saved successfully.")
                    saved_glossary_chunks_count += 1
                except Exception as e:
                    logger.error(f"Error saving glossary chunk {chunk_filename.name}: {e}")
            logger.info(f"Finished saving glossaries into {saved_glossary_chunks_count} chunk file(s).")
    else:
        logger.info("No glossaries were extracted to save.")

    # --- Перемещение очищенных файлов из temp_cleaned_path в cleaned_output_path ---
    logger.info(
        f"Moving cleaned files from temporary directory '{temp_cleaned_path}' to final cleaned output directory '{cleaned_output_path}'...")
    moved_cleaned_files_count = 0
    failed_to_move_count = 0
    for item in temp_cleaned_path.iterdir():  # Итерируемся по содержимому временной папки
        if item.is_file():
            target_path_in_cleaned_output = cleaned_output_path / item.name
            try:
                # shutil.move перезапишет файл в целевой папке, если он там уже существует
                shutil.move(str(item), str(target_path_in_cleaned_output))
                moved_cleaned_files_count += 1
            except Exception as e:
                logger.error(f"Failed to move cleaned file {item.name} to {target_path_in_cleaned_output}: {e}")
                failed_to_move_count += 1
    logger.info(
        f"Finished moving cleaned files. Moved: {moved_cleaned_files_count}, Failed moves: {failed_to_move_count}.")

    # Опциональная очистка временной папки после перемещения
    # (можно закомментировать, если нужно посмотреть содержимое temp_cleaned_path для отладки)
    try:
        # Удаляем только если она пуста или содержит только папки (на случай ошибок)
        if not any(temp_cleaned_path.iterdir()):  # Проверка, пуста ли папка
            shutil.rmtree(temp_cleaned_path)
            logger.info(f"Successfully removed empty temporary directory: {temp_cleaned_path}")
        else:  # Если не пуста после попыток перемещения, возможно, что-то пошло не так
            logger.warning(
                f"Temporary directory {temp_cleaned_path} is not empty after move operation. Manual check advised.")
            # Можно добавить принудительное удаление, если уверены:
            # shutil.rmtree(temp_cleaned_path)
            # logger.info(f"Forcibly removed temporary directory: {temp_cleaned_path}")
    except Exception as e:
        logger.error(f"Could not remove temporary directory {temp_cleaned_path}: {e}")

    logger.info("Glossary extraction and file cleaning process finished.")


# --- HTML Conversion (Остается без изменений, т.к. очистка теперь происходит раньше) ---
# --- Функция build_tome_info ИЗМЕНЕНА на build_volume_info и доработана (Строка 1046 -> 1050) ---
async def build_volume_info(cleaned_files_path: Path) -> Dict[str, Dict[str, Any]]:
    """
    Builds volume information map from cleaned TXT files.
    Returns: {volume_name: {'min_chapter': num, 'order': index, 'chapters': [num, ...]}}
    """
    logger.info(f"Building volume information map from cleaned TXT files in: {cleaned_files_path}")
    volume_data: Dict[str, List[int]] = {}  # {safe_volume_name: [chapter_num, ...]}

    # Сканируем только .txt файлы, так как они содержат информацию о томе в нужном формате
    files_to_scan = list(cleaned_files_path.glob('[0-9][0-9][0-9][0-9]*.txt'))
    if not files_to_scan:
        logger.warning(f"No cleaned TXT files found in {cleaned_files_path} to build volume info.")
        return {}
    logger.debug(f"Scanning {len(files_to_scan)} cleaned TXT files for volume info...")

    # Вспомогательная функция для асинхронного чтения и парсинга каждого файла
    async def scan_txt_file_for_volume(file_path: Path):
        try:
            match = re.match(r'^(\d{4})', file_path.name)
            if not match:
                return  # Пропускаем файлы без корректного номера главы в имени

            chapter_num = int(match.group(1))

            # Очищенные файлы должны иметь следующую структуру для определения тома:
            # Строка 0: Заголовок главы
            # Строка 1: Пустая строка (или отсутствует, если нет тома)
            # Строка 2: Название тома (если есть)
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                # Читаем до 3 строк, т.к. название тома (если есть) обычно на 3-й строке (индекс 2)
                lines = [await f.readline() for _ in range(3)]

            volume_name_raw = "Unknown Volume"  # Имя тома по умолчанию
            if len(lines) >= 3 and lines[1].strip() == "":  # Проверяем, что есть вторая (пустая) и третья строки
                potential_volume_name = lines[2].strip()
                if potential_volume_name:  # Если третья строка не пустая, это название тома
                    volume_name_raw = potential_volume_name
            else:
                logger.debug(
                    f"No volume name found in standard position for {file_path.name}. Assigning to 'Unknown Volume'.")

            # Очищаем имя тома для использования в качестве ключа или имени папки
            safe_volume_name = re.sub(r'[\\/*?:"<>|]', '_', volume_name_raw) if volume_name_raw else "Unknown_Volume"

            if safe_volume_name not in volume_data:
                volume_data[safe_volume_name] = []
            volume_data[safe_volume_name].append(chapter_num)

        except ValueError:  # Ошибка преобразования номера главы в int
            logger.warning(f"Could not parse chapter number for {file_path.name} during volume scan.")
        except Exception as e:
            logger.error(f"Error scanning file {file_path.name} for volume info: {e}", exc_info=False)

    scan_tasks = [asyncio.create_task(scan_txt_file_for_volume(f)) for f in files_to_scan]
    if scan_tasks:
        await asyncio.gather(*scan_tasks)

    if not volume_data:
        logger.warning("No volume information could be extracted from the cleaned files.")
        return {}

    # Сортируем тома по минимальному номеру главы в них для присвоения порядка
    # (volume_min_chapter, safe_volume_name, list_of_chapters_in_volume)
    sorted_volume_meta = []
    for sv_name, chap_list in volume_data.items():
        if chap_list:
            chap_list.sort()  # Сортируем главы внутри каждого тома
            sorted_volume_meta.append((chap_list[0], sv_name, chap_list))

    sorted_volume_meta.sort()  # Сортируем сами тома по их первой главе

    final_volume_info_map: Dict[str, Dict[str, Any]] = {}
    for i, (min_chap, sv_name, chap_list) in enumerate(sorted_volume_meta):
        final_volume_info_map[sv_name] = {
            'min_chapter': min_chap,
            'order': i + 1,  # Порядковый номер тома (1, 2, 3...)
            'chapters': chap_list,  # Список номеров глав в этом томе
            'raw_name': sv_name  # Сохраняем "сырое" безопасное имя для сопоставления
        }

    logger.info(f"Built volume info map for {len(final_volume_info_map)} volumes.")
    if logger.level == logging.DEBUG:  # Логируем карту только если DEBUG уровень
        for vol_name, info in final_volume_info_map.items():
            logger.debug(
                f"Volume: '{vol_name}', Order: {info['order']}, MinChap: {info['min_chapter']}, Chapters: {info['chapters'][:5]}...")  # Первые 5 глав для краткости
    return final_volume_info_map


# --- Функция convert_cleaned_to_html остается без изменений в логике чтения строк, ---
# --- т.к. очистка теперь происходит в extract_glossary_and_clean_files       ---
# --- Использует обновленную build_volume_info (Строка 1078 -> 1133) ---
async def convert_cleaned_to_html(config: Config):
    """Converts cleaned text files to simple HTML files, adding volume titles and using chapter title in filename."""
    cleaned_output_path = Path(config.get('Settings', 'CleanedOutputPath', default='./CleanedOutput'))
    html_output_path = Path(config.get('Settings', 'HtmlOutputPath', default='./HtmlOutput'))

    logger.info(f"Starting HTML conversion from '{cleaned_output_path}' to '{html_output_path}'.")
    if not cleaned_output_path.is_dir():
        logger.error(f"Source path for cleaned files '{cleaned_output_path}' not found. Cannot convert to HTML.")
        return
    html_output_path.mkdir(parents=True, exist_ok=True)

    # Используем обновленную функцию для получения информации о томах
    volume_info_map = await build_volume_info(cleaned_output_path)
    if not volume_info_map:
        logger.warning("Volume information map is empty. HTML files will be generated without H2 volume titles.")

    files_to_convert = sorted(list(cleaned_output_path.glob('[0-9][0-9][0-9][0-9]*.txt')))
    if not files_to_convert:
        logger.info(f"No cleaned files found in {cleaned_output_path} to convert to HTML.")
        return
    logger.info(f"Found {len(files_to_convert)} cleaned files to convert to HTML.")

    tasks = []

    async def process_file_to_html(txt_file_path: Path, vol_info_map: Dict[str, Dict[str, Any]]):
        try:
            chapter_num_from_filename = -1
            match_fn = re.match(r'^(\d{4})', txt_file_path.name)
            if match_fn:
                try:
                    chapter_num_from_filename = int(match_fn.group(1))
                except ValueError:
                    pass

            async with aiofiles.open(txt_file_path, 'r', encoding='utf-8') as infile:
                lines = await infile.readlines()

            if not lines:
                logger.warning(f"Cleaned file {txt_file_path.name} is empty. Skipping HTML conversion.")
                return

            html_parts = []

            # Очищенный файл: первая строка - заголовок главы.
            # Вторая строка - пустая (если есть том).
            # Третья строка - название тома (если есть).
            # Остальное - контент.

            chapter_title_raw = lines[0].strip()
            current_volume_name_raw = None  # "Сырое" название тома из файла
            current_volume_safe_name = None  # Безопасное имя тома
            content_start_index = 1  # Индекс, с которого начинается основной контент главы

            if len(lines) >= 3 and lines[1].strip() == "":  # Есть пустая строка -> возможно, есть том
                potential_volume_name = lines[2].strip()
                if potential_volume_name:
                    current_volume_name_raw = potential_volume_name
                    current_volume_safe_name = re.sub(r'[\\/*?:"<>|]', '_', current_volume_name_raw)
                    content_start_index = 3  # Контент после строки с томом
                    # Пропускаем возможные пустые строки после тома перед контентом
                    while content_start_index < len(lines) and not lines[content_start_index].strip():
                        content_start_index += 1

            # Добавление H2 для тома, если это первая глава тома
            if current_volume_safe_name and current_volume_safe_name in vol_info_map:
                vol_details = vol_info_map[current_volume_safe_name]
                if chapter_num_from_filename != -1 and chapter_num_from_filename == vol_details['min_chapter']:
                    vol_order = vol_details['order']
                    # Используем current_volume_name_raw для отображения, т.к. оно оригинальное
                    html_parts.append(
                        f'<h2 style="text-align: center;">Том {vol_order}. {current_volume_name_raw}</h2>\n')
                    logger.debug(
                        f"Added H2 title for Volume {vol_order} ('{current_volume_name_raw}') in HTML for {txt_file_path.name}")

            # Добавление H3 для главы
            if chapter_num_from_filename != -1:
                html_parts.append(
                    f'<h3 style="text-align: center;">Глава {chapter_num_from_filename}. {chapter_title_raw}</h3>\n')
            else:  # Если номер главы не определен из имени файла
                html_parts.append(f'<h3 style="text-align: center;">{chapter_title_raw}</h3>\n')

            # Обработка основного контента
            for i in range(content_start_index, len(lines)):
                line_content = lines[i].strip()
                if line_content:
                    # Применение базового форматирования Markdown (*italic*, **bold**)
                    processed_line = re.sub(r'\*(.+?)\*', r'<i>\1</i>', line_content)
                    processed_line = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', processed_line)
                    html_parts.append(f"<p>{processed_line}</p>\n")
                else:  # Пустая строка в тексте -> пустая строка или <br> в HTML
                    html_parts.append("\n")  # Или <p>&nbsp;</p> для видимого пустого абзаца

            # Добавление разделителя глав Sigil (если он используется)
            html_parts.append('<hr class="sigil_split_marker" />\n')  # Стандартный разделитель

            # Формирование имени HTML файла
            safe_chapter_title_for_fn = re.sub(r'[\\/*?:"<>|]', '_', chapter_title_raw)
            safe_chapter_title_for_fn = safe_chapter_title_for_fn[:150].strip()  # Ограничение длины имени файла

            html_filename_str = ""
            if chapter_num_from_filename != -1:
                html_filename_str = f"{chapter_num_from_filename:04d} - {safe_chapter_title_for_fn}.html"
            else:
                html_filename_str = f"{safe_chapter_title_for_fn}.html"

            final_html_filepath = html_output_path / html_filename_str
            async with aiofiles.open(final_html_filepath, 'w', encoding='utf-8') as outfile:
                await outfile.writelines(html_parts)
            logger.debug(f"Successfully converted '{txt_file_path.name}' to HTML file '{html_filename_str}'")

        except Exception as e:
            logger.error(f"Error converting file {txt_file_path.name} to HTML: {e}", exc_info=True)

    for txt_file_path_item in files_to_convert:
        tasks.append(asyncio.create_task(process_file_to_html(txt_file_path_item, volume_info_map)))

    if tasks:
        await asyncio.gather(*tasks)

    logger.info(f"HTML conversion finished. Results are in '{html_output_path}'.")


# --- НОВАЯ ФУНКЦИЯ: Конвертация в DOCX (Строка 1161 -> 1239) ---
# (add_formatted_run остается той же)
def add_formatted_run(paragraph, text_segment):
    """Добавляет текст в параграф с распознаванием **bold** и *italic*."""
    parts = re.split(r'(\*\*(?:[^*]|(?<!\*)\*(?!\*))*?\*\*|\*(?:[^*]|(?<!\*)\*(?!\*))*?\*)', text_segment)
    for part in parts:
        if not part: continue
        if part.startswith('**') and part.endswith('**'):
            content = part[2:-2]
            run = paragraph.add_run(content)
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            content = part[1:-1]
            run = paragraph.add_run(content)
            run.italic = True
        else:
            paragraph.add_run(part)


async def convert_cleaned_to_docx(config: Config):
    """Converts cleaned text files to DOCX files, preserving structure and basic formatting."""
    cleaned_output_path = Path(config.get('Settings', 'CleanedOutputPath', default='./CleanedOutput'))
    docx_output_path = Path(config.get('Settings', 'DocxOutputPath', default='./DocxOutput'))

    logger.info(f"Starting DOCX conversion from '{cleaned_output_path}' to '{docx_output_path}'.")
    if not cleaned_output_path.is_dir():
        logger.error(f"Cleaned source path '{cleaned_output_path}' not found. Cannot convert to DOCX.")
        return
    docx_output_path.mkdir(parents=True, exist_ok=True)

    # Используем обновленную build_volume_info
    volume_info_map = await build_volume_info(cleaned_output_path)
    if not volume_info_map:
        logger.warning("Volume information map is empty. DOCX files will be generated without Volume titles.")

    files_to_convert = sorted(list(cleaned_output_path.glob('[0-9][0-9][0-9][0-9]*.txt')))
    if not files_to_convert:
        logger.info(f"No cleaned files found in {cleaned_output_path} to convert to DOCX.")
        return
    logger.info(f"Found {len(files_to_convert)} cleaned files to convert to DOCX.")

    converted_count = 0
    for txt_file_path in files_to_convert:  # Обрабатываем синхронно из-за python-docx
        try:
            chapter_num_from_filename = -1
            match_fn = re.match(r'^(\d{4})', txt_file_path.name)
            if match_fn:
                try:
                    chapter_num_from_filename = int(match_fn.group(1))
                except ValueError:
                    pass

            # Используем синхронное чтение для DOCX части
            with open(txt_file_path, 'r', encoding='utf-8') as infile:
                lines = infile.readlines()

            if not lines:
                logger.warning(f"Cleaned file {txt_file_path.name} is empty. Skipping DOCX conversion.")
                continue

            document = docx.Document()
            # (Можно настроить стили по умолчанию здесь, если нужно)
            # style = document.styles['Normal']
            # font = style.font; font.name = 'Times New Roman'; font.size = Pt(12)

            chapter_title_raw = lines[0].strip()
            current_volume_name_raw = None
            current_volume_safe_name = None
            content_start_index = 1

            if len(lines) >= 3 and lines[1].strip() == "":
                potential_volume_name = lines[2].strip()
                if potential_volume_name:
                    current_volume_name_raw = potential_volume_name
                    current_volume_safe_name = re.sub(r'[\\/*?:"<>|]', '_', current_volume_name_raw)
                    content_start_index = 3
                    while content_start_index < len(lines) and not lines[content_start_index].strip():
                        content_start_index += 1

            if current_volume_safe_name and current_volume_safe_name in volume_info_map:
                vol_details = volume_info_map[current_volume_safe_name]
                if chapter_num_from_filename != -1 and chapter_num_from_filename == vol_details['min_chapter']:
                    vol_order = vol_details['order']
                    h2 = document.add_heading(f"Том {vol_order}. {current_volume_name_raw}", level=2)
                    h2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    logger.debug(f"Added H2 (Том) for '{current_volume_name_raw}' in DOCX for {txt_file_path.name}")

            if chapter_num_from_filename != -1:
                h3 = document.add_heading(f"Глава {chapter_num_from_filename}. {chapter_title_raw}", level=3)
            else:
                h3 = document.add_heading(chapter_title_raw, level=3)
            h3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            for i in range(content_start_index, len(lines)):
                line_content = lines[i].strip()
                if line_content:
                    p = document.add_paragraph()
                    add_formatted_run(p, line_content)  # Используем хелпер для **bold** и *italic*
                else:
                    document.add_paragraph()  # Пустой параграф для разделения

            safe_chapter_title_for_fn = re.sub(r'[\\/*?:"<>|]', '_', chapter_title_raw)
            safe_chapter_title_for_fn = safe_chapter_title_for_fn[:150].strip()

            docx_filename_str = ""
            if chapter_num_from_filename != -1:
                docx_filename_str = f"{chapter_num_from_filename:04d} - {safe_chapter_title_for_fn}.docx"
            else:
                docx_filename_str = f"{safe_chapter_title_for_fn}.docx"

            final_docx_filepath = docx_output_path / docx_filename_str
            document.save(final_docx_filepath)
            logger.debug(f"Successfully converted '{txt_file_path.name}' to DOCX file '{docx_filename_str}'")
            converted_count += 1

        except Exception as e:
            logger.error(f"Error converting file {txt_file_path.name} to DOCX: {e}", exc_info=True)

    logger.info(f"DOCX conversion finished. Converted {converted_count} files. Results are in '{docx_output_path}'.")


# --- Volume Sorting Function (Остается без изменений, т.к. читает из Output) ---
# --- ИЗМЕНЕНИЕ: Учитываем удаление пустых строк при чтении из OutputPath (Строка 1282 -> 1360) ---
# --- Использует обновленную build_volume_info (косвенно, через логику определения тома)
def sort_files_into_volumes(config: Config):
    """Sorts translated files from OutputPath into VolumeSortPath by volume."""
    # Эта функция сортирует "сырые" выходные файлы из OutputPath,
    # а не уже очищенные из CleanedOutputPath.
    # Логика определения тома здесь должна быть самодостаточной.
    output_path = Path(config.get('Settings', 'OutputPath', default='./Output'))
    volume_sort_path = Path(config.get('Settings', 'VolumeSortPath', default='./Volumes'))
    logger.info(f"Starting volume sorting from RAW output '{output_path}' to '{volume_sort_path}'.")

    if not output_path.is_dir():
        logger.error(f"Raw output path '{output_path}' not found. Cannot sort files into volumes.")
        return
    volume_sort_path.mkdir(parents=True, exist_ok=True)

    # {safe_volume_name: [(chapter_num, safe_chapter_title, original_filepath, raw_volume_name_from_file), ...]}
    volume_chapters_map: Dict[str, List[Tuple[int, str, Path, str]]] = {}
    processed_files_count = 0

    # Ищем .txt файлы в OutputPath
    for original_file_path in output_path.glob('[0-9][0-9][0-9][0-9]*.txt'):
        chapter_num_from_filename = -1
        try:
            chapter_num_from_filename = int(original_file_path.name[:4])
        except ValueError:
            logger.warning(
                f"Could not parse chapter number from filename {original_file_path.name}. Skipping sort for this file.")
            continue

        try:
            with open(original_file_path, "r", encoding="utf-8", errors='replace') as f:
                lines = f.readlines()

            if not lines:
                logger.warning(f"File {original_file_path.name} is empty. Skipping sort for this file.");
                continue

            # Логика извлечения заголовка главы и тома из "сырого" файла
            # (может содержать разделитель глоссария и сам глоссарий)
            # Сначала удаляем секцию глоссария, если она есть
            content_str = "".join(lines)
            text_before_glossary = content_str
            if GLOSSARY_SEPARATOR in content_str:
                text_before_glossary = content_str.split(GLOSSARY_SEPARATOR, 1)[0]

            cleaned_lines = text_before_glossary.splitlines()

            # Ищем первую непустую строку для заголовка главы
            first_content_line_idx = 0
            while first_content_line_idx < len(cleaned_lines) and not cleaned_lines[first_content_line_idx].strip():
                first_content_line_idx += 1

            if first_content_line_idx >= len(cleaned_lines):
                logger.warning(
                    f"{original_file_path.name} effectively empty after removing glossary and leading blanks. Skipping.");
                continue

            chapter_title_raw = cleaned_lines[first_content_line_idx].strip()

            # Ищем название тома (обычно через 1 пустую строку после заголовка главы)
            volume_line_idx = first_content_line_idx + 2
            volume_name_raw_from_file = "Unknown Volume"  # Имя тома по умолчанию
            if volume_line_idx < len(cleaned_lines) and \
                    (first_content_line_idx + 1) < len(cleaned_lines) and \
                    cleaned_lines[first_content_line_idx + 1].strip() == "":  # Проверяем пустую строку
                potential_volume_name = cleaned_lines[volume_line_idx].strip()
                if potential_volume_name:
                    volume_name_raw_from_file = potential_volume_name

            safe_chapter_title_for_fn = re.sub(r'[\\/*?:"<>|]', '_', chapter_title_raw)
            safe_volume_name_key = re.sub(r'[\\/*?:"<>|]', '_',
                                          volume_name_raw_from_file) if volume_name_raw_from_file else "Unknown_Volume"

            if safe_volume_name_key not in volume_chapters_map:
                volume_chapters_map[safe_volume_name_key] = []

            volume_chapters_map[safe_volume_name_key].append(
                (chapter_num_from_filename, safe_chapter_title_for_fn, original_file_path, volume_name_raw_from_file)
            )
            processed_files_count += 1

        except Exception as e:
            logger.error(f"Error reading/parsing file {original_file_path.name} for volume sorting: {e}",
                         exc_info=False)

    if not volume_chapters_map:
        logger.info(f"No valid chapter files found in '{output_path}' to sort into volumes.")
        return
    logger.info(
        f"Identified {len(volume_chapters_map)} potential volumes from {processed_files_count} files in '{output_path}'.")

    # Сортируем тома по номеру первой главы в них
    # (min_chapter_in_volume, safe_volume_name_key, raw_volume_name_for_display)
    # Для raw_volume_name_for_display берем из первой главы тома, предполагая, что оно там консистентно.
    volume_order_metadata = []
    for sv_name, chap_data_list in volume_chapters_map.items():
        if chap_data_list:
            min_chap_num = min(cd[0] for cd in chap_data_list)
            # Берем "сырое" имя тома из данных первой главы (по номеру) этого тома
            # Сначала отсортируем chap_data_list по номеру главы, чтобы взять правильное "сырое" имя
            chap_data_list.sort(key=lambda x: x[0])
            raw_name_display = chap_data_list[0][3] if chap_data_list[0][
                3] else sv_name  # Используем сырое имя, если есть
            volume_order_metadata.append((min_chap_num, sv_name, raw_name_display))

    volume_order_metadata.sort()  # Сортируем сами тома

    # Создание папок томов и копирование/переименование файлов
    for vol_idx, (min_chap, sv_name_key, raw_vol_name_disp) in enumerate(volume_order_metadata):
        volume_order_num = vol_idx + 1  # Порядковый номер тома 1, 2, ...
        # Формируем имя папки тома: "0001_БезопасноеИмяТома"
        numbered_volume_dirname = f"{volume_order_num:04d}_{sv_name_key}"
        target_volume_dir_path = volume_sort_path / numbered_volume_dirname
        target_volume_dir_path.mkdir(exist_ok=True)

        chapters_for_this_volume = sorted(volume_chapters_map[sv_name_key],
                                          key=lambda x: x[0])  # Сортируем главы внутри тома

        for chap_num, safe_chap_title, orig_fp, _ in chapters_for_this_volume:
            # Новое имя файла: "0001 - БезопасноеИмяГлавы.txt"
            new_target_filename = f"{chap_num:04d} - {safe_chap_title}.txt"
            new_target_filepath = target_volume_dir_path / new_target_filename
            try:
                # Копируем исходный файл из OutputPath в папку тома с новым именем
                # Содержимое файла не меняем на этом этапе, просто копируем как есть.
                # Очистка и добавление заголовка тома происходят в других шагах (extract_glossary, convert_to_html/docx)
                shutil.copy2(orig_fp, new_target_filepath)
                logger.debug(f"Sorted '{orig_fp.name}' to '{new_target_filepath}'")
            except Exception as e:
                logger.error(f"Error sorting file {orig_fp.name} to {new_target_filepath}: {e}")

    logger.info(f"Finished sorting files from '{output_path}' into volume subdirectories in '{volume_sort_path}'.")


# --- START OF NEW FUNCTION find_chapters_without_glossary_marker (Строка 1343 -> 1497) ---
async def find_chapters_without_glossary_marker(config: Config):
    """
    Scans files in the OutputPath and logs the chapter numbers
    where the GLOSSARY_SEPARATOR was not found.
    """
    output_path = Path(config.get('Settings', 'OutputPath', default='./Output'))
    logger.info(f"Scanning for missing glossary markers ('{GLOSSARY_SEPARATOR}') in path: {output_path}")

    if not output_path.is_dir():
        logger.error(f"Output path '{output_path}' not found or is not a directory. Cannot scan.")
        return

    files_to_scan = sorted(list(output_path.glob('[0-9][0-9][0-9][0-9]*.txt')))
    if not files_to_scan:
        logger.info("No chapter files found in the output path to scan.")
        return

    logger.info(f"Found {len(files_to_scan)} files to scan...")
    chapters_without_marker = []
    checked_files_count = 0

    async def check_file(file_path: Path):
        nonlocal checked_files_count  # Разрешаем изменять внешнюю переменную
        chapter_num = -1
        try:
            match = re.match(r'^(\d{4})', file_path.name)
            if match:
                chapter_num = int(match.group(1))
            else:
                logger.warning(f"Could not parse chapter number from filename: {file_path.name}. Skipping check.")
                return

            async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='replace') as infile:
                content = await infile.read()

            if GLOSSARY_SEPARATOR not in content:
                logger.debug(f"Glossary marker NOT found in chapter {chapter_num} ({file_path.name})")
                chapters_without_marker.append(chapter_num)
            else:
                logger.debug(f"Glossary marker found in chapter {chapter_num} ({file_path.name})")

            # Важно: увеличивать счетчик здесь, чтобы он отражал количество фактически проверенных файлов
            # Если использовать `nonlocal` для chapters_without_marker, то для checked_files_count тоже нужно
            # checked_files_count +=1 # Это вызовет UnboundLocalError без nonlocal
            # Вместо этого, пусть check_file возвращает результат, а собираем его снаружи

        except ValueError:  # Ошибка при int(match.group(1))
            logger.warning(f"Could not parse chapter number as integer from {file_path.name}. Skipping.")
        except Exception as e:
            logger.error(f"Error processing file {file_path.name} during marker check: {e}", exc_info=False)
        return chapter_num, GLOSSARY_SEPARATOR not in content if chapter_num != -1 else None  # Возвращаем результат

    # Создаем и запускаем задачи для проверки файлов
    # tasks = [asyncio.create_task(check_file(f)) for f in files_to_scan]
    # if tasks:
    #     await asyncio.gather(*tasks)

    # Альтернативный сбор результатов для корректного инкремента checked_files_count
    temp_chapters_without_marker = []
    for f_path in files_to_scan:
        res = await check_file(f_path)
        checked_files_count += 1  # Считаем только те, что были переданы в check_file
        if res:
            chap_num, is_missing = res
            if chap_num is not None and is_missing:
                temp_chapters_without_marker.append(chap_num)

    chapters_without_marker = sorted(list(set(temp_chapters_without_marker)))  # Уникальные и отсортированные

    logger.info(f"Scan complete. Checked {checked_files_count} files.")
    if chapters_without_marker:
        logger.warning(f"Found {len(chapters_without_marker)} chapters without the glossary marker:")
        logger.warning(f"Chapters missing marker: {chapters_without_marker}")
    else:
        logger.info("Glossary marker found in all checked chapter files (or no files matched the pattern).")


# --- START OF NEW FUNCTION merge_cleaned_files (Строка 1560) ---
async def _merge_txt_files(file_paths: List[Path], output_filepath: Path):
    """Helper to merge multiple TXT files into one."""
    logger.debug(f"Merging {len(file_paths)} TXT files into {output_filepath}")
    try:
        first_file = True
        async with aiofiles.open(output_filepath, 'wb') as outfile:  # Open in binary append mode
            for file_path in file_paths:
                if not first_file:
                    # Добавляем простой разделитель между файлами, если это не первый файл
                    # Можно настроить или убрать, если структура файлов уже это подразумевает
                    await outfile.write(b"\n\n-----\n\n")  # Бинарный разделитель
                first_file = False

                try:
                    async with aiofiles.open(file_path, 'rb') as infile:
                        while True:
                            chunk = await infile.read(8192)  # Читаем по 8KB
                            if not chunk:
                                break
                            await outfile.write(chunk)
                except FileNotFoundError:
                    logger.warning(f"TXT file not found during merge: {file_path}. Skipping.")
                except Exception as e:
                    logger.error(f"Error reading TXT file {file_path} during merge: {e}")
        logger.info(f"Successfully merged TXT files into {output_filepath}")
    except Exception as e:
        logger.error(f"Failed to merge TXT files into {output_filepath}: {e}")


async def _merge_html_files(file_paths: List[Path], output_filepath: Path):
    """Helper to merge multiple HTML files into one."""
    logger.debug(f"Merging {len(file_paths)} HTML files into {output_filepath}")

    # Базовый HTML шаблон для объединенного файла
    # Можно добавить стили в <head> при необходимости
    html_shell_start = """<!DOCTYPE html>
<html lang="ru">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Объединенный HTML Документ</title>
    <style>
        body { font-family: sans-serif; margin: 20px; line-height: 1.6; }
        h2, h3 { text-align: center; }
        .sigil_split_marker { margin-top: 2em; margin-bottom: 2em; }
        /* Дополнительные стили можно добавить здесь */
    </style>
</head>
<body>
"""
    html_shell_end = """
</body>
</html>
"""
    try:
        async with aiofiles.open(output_filepath, 'w', encoding='utf-8') as outfile:
            await outfile.write(html_shell_start)
            for i, file_path in enumerate(file_paths):
                try:
                    async with aiofiles.open(file_path, 'r', encoding='utf-8') as infile:
                        content = await infile.read()
                        # Исходные HTML файлы уже содержат <hr class="sigil_split_marker" />
                        # Просто добавляем их содержимое.
                        await outfile.write(content)
                        if i < len(file_paths) - 1:  # Добавляем дополнительный разрыв, если это не последний файл
                            await outfile.write("\n<hr />\n")  # Явный HR между контентом файлов

                except FileNotFoundError:
                    logger.warning(f"HTML file not found during merge: {file_path}. Skipping.")
                except Exception as e:
                    logger.error(f"Error reading HTML file {file_path} during merge: {e}")
            await outfile.write(html_shell_end)
        logger.info(f"Successfully merged HTML files into {output_filepath}")
    except Exception as e:
        logger.error(f"Failed to merge HTML files into {output_filepath}: {e}")


async def _merge_docx_files(file_paths: List[Path], output_filepath: Path):
    """Helper to merge multiple DOCX files into one."""
    logger.debug(f"Merging {len(file_paths)} DOCX files into {output_filepath}")

    if not file_paths:
        logger.warning("No DOCX files provided to merge.")
        return

    # Создаем первый документ, который станет основой
    try:
        merged_document = docx.Document(file_paths[0])  # Используем первый файл как основу
        # Добавляем разрыв страницы в конце первого документа, перед добавлением следующих
        # merged_document.add_page_break() # Не всегда нужно после первого, если он уже заканчивается как надо.
    except FileNotFoundError:
        logger.error(f"Base DOCX file {file_paths[0]} not found. Cannot start merge.")
        return
    except Exception as e:
        logger.error(f"Error opening base DOCX file {file_paths[0]}: {e}")
        return

    # Последовательно добавляем остальные документы
    for i in range(1, len(file_paths)):
        file_path = file_paths[i]
        try:
            source_doc = docx.Document(file_path)
            # Добавляем разрыв страницы перед каждым новым документом (кроме первого)
            merged_document.add_page_break()

            for element in source_doc.element.body:
                # Копируем элементы из тела исходного документа в объединенный
                # Этот метод сохраняет большинство форматирования.
                if isinstance(element, (CT_P, CT_Tbl)):  # Копируем параграфы и таблицы
                    merged_document.element.body.append(element)
                # Можно добавить обработку других типов элементов (CT_SectPr для свойств секций),
                # но это усложнит код. Для простого слияния этого обычно достаточно.

        except FileNotFoundError:
            logger.warning(f"DOCX file not found during merge: {file_path}. Skipping.")
        except Exception as e:
            logger.error(f"Error reading/appending DOCX file {file_path}: {e}")
            # Можно решить, продолжать ли слияние или остановить при ошибке

    try:
        # Сохранение объединенного документа (синхронная операция)
        await asyncio.to_thread(merged_document.save, output_filepath)
        logger.info(f"Successfully merged DOCX files into {output_filepath}")
    except Exception as e:
        logger.error(f"Failed to save merged DOCX file {output_filepath}: {e}")


async def merge_cleaned_files(config: Config):
    """
    Merges cleaned files (txt, html, docx) based on settings in config.yml -> MergeSettings.
    """
    logger.info("Starting process to merge cleaned files.")

    merge_config = config.get('MergeSettings')
    if not merge_config:
        logger.warning("MergeSettings section not found in config. Skipping merge process.")
        return

    output_root_path = Path(merge_config.get('OutputPath', './MergedOutput'))
    files_per_chunk_setting = merge_config.get('FilesToMergePerChunk', 0)
    merge_by_volume_setting = merge_config.get('MergeByVolume', False)
    start_chapter_num_filter = merge_config.get('StartChapterNumber', 1)

    path_for_volume_def_str = merge_config.get('PathForVolumeDefinition', './CleanedOutput')
    path_for_volume_def = Path(path_for_volume_def_str)

    type_settings = merge_config.get('Types', {})
    if not type_settings:
        logger.warning("No file types configured for merging in MergeSettings.Types. Skipping.")
        return

    output_root_path.mkdir(parents=True, exist_ok=True)

    volume_info_map = None
    if merge_by_volume_setting:
        logger.info(f"MergeByVolume is enabled. Attempting to build volume info from: {path_for_volume_def}")
        volume_info_map = await build_volume_info(path_for_volume_def)  # build_volume_info уже было доработано
        if not volume_info_map:
            logger.warning(
                "Could not build volume map. MergeByVolume cannot proceed effectively. Merging might be incorrect or skipped.")

    for file_ext, type_config in type_settings.items():
        if not type_config.get('Enabled', False):
            logger.info(f"Merging for type '{file_ext}' is disabled. Skipping.")
            continue

        source_path_str = type_config.get('SourcePath')
        if not source_path_str:
            logger.warning(f"SourcePath not configured for type '{file_ext}'. Skipping.")
            continue

        source_path = Path(source_path_str)
        if not source_path.is_dir():
            logger.warning(f"SourcePath '{source_path}' for type '{file_ext}' is not a valid directory. Skipping.")
            continue

        logger.info(f"Processing merge for type: '{file_ext}' from source: '{source_path}'")

        all_files_of_type = sorted(list(source_path.glob(f"*.{file_ext}")))

        eligible_files_map: Dict[int, Path] = {}
        for f_path in all_files_of_type:
            match = re.match(r'^(\d{4})', f_path.name)
            if match:
                try:
                    chap_num = int(match.group(1))
                    if chap_num >= start_chapter_num_filter:
                        eligible_files_map[chap_num] = f_path
                except ValueError:
                    logger.debug(
                        f"Could not parse chapter number from {f_path.name} for type {file_ext}. Skipping for start_chapter filter.")

        if not eligible_files_map:
            logger.info(
                f"No eligible files found for type '{file_ext}' (after StartChapterNumber filter) in '{source_path}'.")
            continue

        eligible_file_paths = [eligible_files_map[cn] for cn in sorted(eligible_files_map.keys())]

        if merge_by_volume_setting and volume_info_map:
            logger.info(f"Merging type '{file_ext}' by volume.")
            for vol_safe_name, vol_details in volume_info_map.items():
                vol_order = vol_details['order']
                # vol_chapters_list содержит ВСЕ теоретические главы тома из build_volume_info
                vol_chapters_list_from_map = vol_details['chapters']

                files_for_this_volume: List[Path] = []
                # Собираем ФАКТИЧЕСКИЕ номера глав, которые войдут в слияние для этого тома (после всех фильтров)
                actual_chapter_numbers_in_volume_merge: List[int] = []

                for chap_num_in_vol_map in vol_chapters_list_from_map:
                    # Проверяем, что глава есть в eligible_files_map (т.е. она существует и прошла start_chapter_num_filter)
                    if chap_num_in_vol_map in eligible_files_map:
                        files_for_this_volume.append(eligible_files_map[chap_num_in_vol_map])
                        actual_chapter_numbers_in_volume_merge.append(chap_num_in_vol_map)

                if not files_for_this_volume:  # или not actual_chapter_numbers_in_volume_merge
                    logger.debug(
                        f"No eligible files of type '{file_ext}' found for volume '{vol_safe_name}' (Order {vol_order}) after filtering. Skipping merge for this volume/type.")
                    continue

                # Сортируем фактические номера глав на случай, если vol_chapters_list_from_map был не отсортирован
                actual_chapter_numbers_in_volume_merge.sort()
                first_chapter_in_merge_num = actual_chapter_numbers_in_volume_merge[0]
                last_chapter_in_merge_num = actual_chapter_numbers_in_volume_merge[-1]

                chapters_range_str = f"Chapters_{first_chapter_in_merge_num:04d}-{last_chapter_in_merge_num:04d}"

                display_vol_name = vol_details.get('raw_name', vol_safe_name)
                fn_safe_vol_name = re.sub(r'[^\w\s-]', '', display_vol_name).strip().replace(' ', '_')
                fn_safe_vol_name = fn_safe_vol_name[:50]

                # --- ИЗМЕНЕННАЯ СТРОКА ДЛЯ ИМЕНИ ФАЙЛА ---
                output_filename = f"Merged_Volume_{vol_order:02d}_{fn_safe_vol_name}_{chapters_range_str}.{file_ext}"
                # --- КОНЕЦ ИЗМЕНЕНИЯ ---
                output_filepath_for_volume = output_root_path / output_filename

                logger.info(
                    f"Merging {len(files_for_this_volume)} files for Volume {vol_order} ('{display_vol_name}', chapters {first_chapter_in_merge_num:04d}-{last_chapter_in_merge_num:04d}) into {output_filepath_for_volume}")
                if file_ext == 'txt':
                    await _merge_txt_files(files_for_this_volume, output_filepath_for_volume)
                elif file_ext == 'html':
                    await _merge_html_files(files_for_this_volume, output_filepath_for_volume)
                elif file_ext == 'docx':
                    # Передаем отсортированный список путей к файлам, соответствующих actual_chapter_numbers_in_volume_merge
                    # files_for_this_volume уже должен быть в правильном порядке, если eligible_files_map.keys() был отсортирован
                    # и vol_chapters_list_from_map тоже. Для надежности можно пересортировать files_for_this_volume
                    # по номеру главы из имени файла.
                    files_for_this_volume.sort(
                        key=lambda p: int(re.match(r'^(\d{4})', p.name).group(1)) if re.match(r'^(\d{4})',
                                                                                              p.name) else 0)
                    await _merge_docx_files(files_for_this_volume, output_filepath_for_volume)

        elif not merge_by_volume_setting:
            logger.info(f"Merging type '{file_ext}' by chunks or all-in-one.")
            if not eligible_file_paths:
                logger.info(f"No eligible files to merge for type '{file_ext}' after all filters.")
                continue

            num_total_eligible_files = len(eligible_file_paths)

            chunk_size = files_per_chunk_setting
            if chunk_size <= 0:
                chunk_size = num_total_eligible_files
            if chunk_size == 0 and num_total_eligible_files > 0:
                chunk_size = num_total_eligible_files

            for i in range(0, num_total_eligible_files, chunk_size):
                current_chunk_paths = eligible_file_paths[i: i + chunk_size]
                if not current_chunk_paths: continue

                first_chap_match = re.match(r'^(\d{4})', current_chunk_paths[0].name)
                last_chap_match = re.match(r'^(\d{4})', current_chunk_paths[-1].name)

                start_c = first_chap_match.group(1) if first_chap_match else "UnknownStart"
                end_c = last_chap_match.group(1) if last_chap_match else "UnknownEnd"
                chunk_num_display = (i // chunk_size) + 1

                output_filename_chunk = ""
                if num_total_eligible_files == chunk_size:
                    output_filename_chunk = f"Merged_All_Chapters_{start_c}-{end_c}.{file_ext}"
                else:
                    output_filename_chunk = f"Merged_Chunk_{chunk_num_display:03d}_Chapters_{start_c}-{end_c}.{file_ext}"

                output_filepath_for_chunk = output_root_path / output_filename_chunk

                logger.info(
                    f"Merging chunk {chunk_num_display} ({len(current_chunk_paths)} files, chapters {start_c}-{end_c}) for type '{file_ext}' into {output_filepath_for_chunk}")

                if file_ext == 'txt':
                    await _merge_txt_files(current_chunk_paths, output_filepath_for_chunk)
                elif file_ext == 'html':
                    await _merge_html_files(current_chunk_paths, output_filepath_for_chunk)
                elif file_ext == 'docx':
                    await _merge_docx_files(current_chunk_paths, output_filepath_for_chunk)
        else:
            logger.warning(
                f"Skipping merge for type '{file_ext}' because MergeByVolume is true but volume information is unavailable.")

    logger.info("Finished merging cleaned files process.")



# --- Main Execution ---
# --- ИЗМЕНЕНИЕ: Добавлен новый режим `find_missing_glossary` (Строка 1406 -> 1904) ---
# --- ИЗМЕНЕНИЕ: Добавлен новый режим `merge_cleaned` (Строка 1904 -> 1906) ---
if __name__ == "__main__":
    try:
        config = Config(CONFIG_PATH)
        run_mode = config.get('Settings', 'RunMode', default='async').lower()
        logger.info(f"Selected RunMode: {run_mode}")

        # Запуск asyncio.run для асинхронных функций, прямой вызов для синхронных
        if run_mode == 'async':
            asyncio.run(main_async(config))
        elif run_mode == 'sequential':
            # main_sequential содержит вызовы asyncio.run для process_single_file,
            # но сама она не async. Для единообразия можно и ее обернуть, но пока так.
            # Чтобы использовать await внутри main_sequential, ее нужно сделать async
            # Для простоты оставим как есть, но это означает, что api_call_delay внутри нее будет time.sleep
            # ИСПРАВЛЕНО: main_sequential теперь использует await asyncio.sleep и await process_single_file
            asyncio.run(main_sequential(config))  # Теперь main_sequential тоже может быть async
        elif run_mode == 'sort':
            sort_files_into_volumes(config)  # Синхронная
        elif run_mode == 'extract_glossary':
            asyncio.run(extract_glossary_and_clean_files(config))
        elif run_mode == 'convert_to_html':
            asyncio.run(convert_cleaned_to_html(config))
        elif run_mode == 'convert_to_docx':
            asyncio.run(convert_cleaned_to_docx(config))
        elif run_mode == 'find_missing_glossary':  # Проверка маркеров глоссария
            asyncio.run(find_chapters_without_glossary_marker(config))
        # --- НОВЫЙ РЕЖИМ ---
        elif run_mode == 'merge_cleaned':
            asyncio.run(merge_cleaned_files(config))
        else:
            logger.error(
                f"Invalid RunMode '{run_mode}'. Available: async, sequential, sort, "
                f"extract_glossary, convert_to_html, convert_to_docx, "
                f"find_missing_glossary, merge_cleaned"
            )

    except SystemExit as exit_e:
        logger.critical(f"Script exited with SystemExit: {exit_e}")
    except KeyboardInterrupt:
        logger.info("Script interrupted by user (KeyboardInterrupt).")
    except Exception as e:
        logger.critical(f"An unexpected critical error occurred at the top level: {e}", exc_info=True)

    logger.info("Script finished execution.")